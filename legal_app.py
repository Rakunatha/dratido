"""
rdxper-legal — Standalone AI Legal Drafting app
────────────────────────────────────────────────
Split out of the original rdxper.py monolith so the legal-drafting
feature can be deployed and scaled independently of the research-paper
generator.

Pipeline:
  1. Groq API (FREE)  → drafts the full legal document as plain text
  2. python-docx      → formats it into a polished, watermarked .docx

AI Provider:
  Groq (free tier) — https://console.groq.com
  set GROQ_API_KEY=your_key_here

Usage:
  python legal_app.py
"""

import os, uuid, time, secrets, re, json, sqlite3
import urllib.request, urllib.parse
from datetime import datetime
from flask import Flask, request, jsonify, send_file, Response
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import xml.sax.saxutils as _sax

app = Flask(__name__)
app.secret_key = secrets.token_hex(32)

sessions = {}   # token -> {email, user_id, name}
jobs     = {}   # job_id -> {file_path, topic, owner_email}

ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', '')

# ── SQLite DB (users + sessions only — no papers/payments needed here) ────────
DB_PATH = os.environ.get('DB_PATH', 'rdxper_legal.db')

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as db:
        db.executescript("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY, email TEXT UNIQUE NOT NULL,
                name TEXT,
                created_at TEXT DEFAULT (datetime('now')),
                last_login TEXT
            );
            CREATE TABLE IF NOT EXISTS sessions (
                token TEXT PRIMARY KEY,
                email TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now'))
            );
        """)

init_db()
os.makedirs('generated', exist_ok=True)


def session_set(token: str, email: str):
    sessions[token] = {'email': email}
    try:
        with get_db() as db:
            db.execute('INSERT OR REPLACE INTO sessions (token, email) VALUES (?, ?)', (token, email))
    except Exception as e:
        print(f'[session_set] DB error: {e}')


def session_get(token: str):
    if not token:
        return None
    if token in sessions:
        return sessions[token]
    try:
        with get_db() as db:
            row = db.execute('SELECT email FROM sessions WHERE token=?', (token,)).fetchone()
            if row:
                email = row['email']
                user = db.execute('SELECT id, name FROM users WHERE email=?', (email,)).fetchone()
                sessions[token] = {
                    'email': email,
                    'user_id': user['id'] if user else email,
                    'name': user['name'] if user else '',
                }
                return sessions[token]
    except Exception as e:
        print(f'[session_get] DB error: {e}')
    return None


def session_delete(token: str):
    sessions.pop(token, None)
    try:
        with get_db() as db:
            db.execute('DELETE FROM sessions WHERE token=?', (token,))
    except Exception as e:
        print(f'[session_delete] DB error: {e}')


def is_admin(sess: dict) -> bool:
    return bool(sess) and bool(ADMIN_EMAIL) and sess.get('email', '').strip().lower() == ADMIN_EMAIL.strip().lower()


# ═══════════════════════════════════════════════════════════════════════════════
#  AI CLIENT  (Groq — fast free inference)
# ═══════════════════════════════════════════════════════════════════════════════

_GROQ_PREFERRED_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
]


def _get_groq_models(api_key, requests_module):
    headers = {"Authorization": f"Bearer {api_key}"}
    try:
        resp = requests_module.get(
            "https://api.groq.com/openai/v1/models",
            headers=headers,
            timeout=20,
        )
        if resp.status_code != 200:
            return [], f"HTTP {resp.status_code} from Groq /models: {resp.text[:300]}"
        data = resp.json()
        models = data.get("data", []) if isinstance(data, dict) else []
        ids = []
        for item in models:
            if not isinstance(item, dict):
                continue
            model_id = item.get("id")
            if model_id and item.get("active", True):
                ids.append(model_id)
        return ids, None
    except Exception as e:
        return [], f"Could not query Groq /models: {e}"


def _select_groq_models(api_key, requests_module):
    preferred_override = os.environ.get("GROQ_MODEL", "").strip()
    available, discovery_error = _get_groq_models(api_key, requests_module)
    available_set = set(available)

    if preferred_override:
        selected = [preferred_override]
        selected.extend(m for m in _GROQ_PREFERRED_MODELS
                        if m != preferred_override and m in available_set)
    elif available:
        selected = [m for m in _GROQ_PREFERRED_MODELS if m in available_set]
        excluded = ("whisper", "guard", "safeguard", "compound")
        selected.extend(
            m for m in available
            if m not in selected and not any(x in m.lower() for x in excluded)
        )
    else:
        selected = [preferred_override] if preferred_override else list(_GROQ_PREFERRED_MODELS)

    return list(dict.fromkeys(selected)), discovery_error


def ai_generate(prompt: str, system: str = "", temperature: float = 0.7) -> str:
    """Call Groq API with requests library + exponential backoff on 429."""
    import requests as _req

    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("GROQ_API_KEY not set. Get a free key at https://console.groq.com")

    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type":  "application/json",
    }

    _GROQ_MODELS, discovery_error = _select_groq_models(api_key, _req)
    print(f"[Groq] Models selected for this key/project: {_GROQ_MODELS}")
    if discovery_error:
        print(f"[Groq] Model discovery warning: {discovery_error}")
    if not _GROQ_MODELS:
        raise RuntimeError(
            "No active Groq text models are available to this API key/project. "
            "Check Groq Project > Settings > Limits/Model Permissions and API key."
        )

    last_error = None

    for model in _GROQ_MODELS:
        payload = {
            "model":       model,
            "messages":    messages,
            "temperature": temperature,
            "max_completion_tokens": 4096,
            "stream":      False,
        }

        for attempt in range(3):
            try:
                resp = _req.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=90,
                )
            except _req.exceptions.Timeout:
                last_error = f"Timeout on {model}"
                print(f"[Groq] Timeout on {model}, trying next...")
                break
            except _req.exceptions.RequestException as e:
                last_error = f"Request error on {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            status = resp.status_code

            if status == 429:
                wait = 2 ** (attempt + 2)
                last_error = f"429 rate-limited on {model} (attempt {attempt+1})"
                print(f"[Groq] 429 on {model}, waiting {wait}s...")
                time.sleep(wait)
                continue

            if status in (400, 402, 404, 503):
                body = resp.text[:300]
                last_error = f"HTTP {status} on {model}: {body}"
                print(f"[Groq] {status} on {model} (skipping): {body[:120]}")
                break

            if status != 200:
                last_error = f"HTTP {status} on {model}: {resp.text[:300]}"
                print(f"[Groq] Unexpected {status} on {model}: {resp.text[:120]}")
                break

            try:
                data = resp.json()
            except Exception as e:
                last_error = f"JSON parse error on {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            if "error" in data:
                err = data["error"]
                last_error = f"API error on {model}: {err}"
                print(f"[Groq] {last_error}")
                err_str = str(err).lower()
                if "rate" in err_str or "quota" in err_str or "limit" in err_str:
                    wait = 2 ** (attempt + 2)
                    print(f"[Groq] Quota error, waiting {wait}s...")
                    time.sleep(wait)
                    continue
                break

            try:
                text = (data["choices"][0]["message"]["content"] or "").strip()
            except (KeyError, IndexError, TypeError) as e:
                last_error = f"Unexpected shape from {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            if not text:
                last_error = f"Empty content from {model}"
                print(f"[Groq] {last_error}")
                break

            print(f"[Groq] \u2713 {model} ({len(text)} chars)")
            return text

        time.sleep(1)

    raise RuntimeError(
        f"All accessible Groq models failed. Last error: {last_error}. "
        "The application queried Groq /models first, so this error now reflects "
        "models visible to your current API key/project. Check Groq Project "
        "model permissions and API key at https://console.groq.com"
    )


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX BUILDING
# ═══════════════════════════════════════════════════════════════════════════════

RDXPER_WATERMARK_TEXT = 'RDXper Legal - A Rakunatha Khrishanth Manathra Creation'


def add_watermark(doc, text: str = RDXPER_WATERMARK_TEXT):
    """Insert a diagonal, semi-transparent watermark into the header of every
    section (the classic Word VML watermark technique), plus a small text
    credit line in the footer as a reliable fallback for viewers that don't
    render VML shapes."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN

    safe_text = _sax.escape(text, {'"': '&quot;'})

    watermark_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<v:shapetype id="_x0000_t136" coordsize="1600,21600" o:spt="136" adj="10800" '
        'path="m@7,0l@8,0m@5,21600l@6,21600e">'
        '<v:formulas>'
        '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 #0"/>'
        '<v:f eqn="sum 0 0 #1"/><v:f eqn="prod #1 2 1"/><v:f eqn="sum 21600 0 #1"/>'
        '<v:f eqn="if #0 #3 0"/><v:f eqn="if #0 21600 #1"/><v:f eqn="if #3 21600 #2"/>'
        '<v:f eqn="if #3 #1 21600"/><v:f eqn="mid #4 #5"/><v:f eqn="mid #6 #7"/><v:f eqn="val #0"/>'
        '</v:formulas>'
        '<v:path textpathok="t" o:connecttype="custom" '
        'o:connectlocs="@9,0;@10,10800;@9,21600;@8,10800" o:connectangles="270,180,90,0"/>'
        '<v:textpath on="t" fitshape="t"/>'
        '</v:shapetype>'
        '<v:shape id="RDXperWatermark" o:spid="_x0000_s2049" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:520pt;height:110pt;'
        'rotation:315;z-index:-251654144;mso-position-horizontal:center;'
        'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
        'mso-position-vertical-relative:margin" o:allowincell="f" fillcolor="#D8D8D8" stroked="f">'
        '<v:fill opacity=".45"/>'
        f'<v:textpath style="font-family:\'Calibri\';font-size:1pt" string="{safe_text}"/>'
        '</v:shape>'
        '</w:pict>'
    )

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ''
        h_para.alignment = _ALIGN.CENTER
        run = h_para.add_run()
        r_el = run._r
        pict = parse_xml(watermark_xml)
        r_el.append(pict)

        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ''
        f_para.alignment = _ALIGN.CENTER
        f_run = f_para.add_run(text)
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
        f_run.italic = True


def build_ai_legal_docx(doc_type: str, ai_text: str) -> str:
    """Convert the AI-drafted plain-text legal document into a formatted,
    watermarked .docx file."""
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    TNR = 'Times New Roman'
    lines = [ln.rstrip() for ln in ai_text.strip().split('\n')]

    numbered_re   = re.compile(r'^\s*(\d{1,3})[\.\)]\s+(.*)$')
    title_written = False

    for ln in lines:
        stripped = ln.strip()
        if not stripped:
            continue
        clean = stripped.strip('#').strip()
        clean = re.sub(r'^\*\*(.*)\*\*$', r'\1', clean).strip()
        clean = clean.lstrip('*').strip()
        if not clean:
            continue

        m = numbered_re.match(clean)
        if not title_written and not m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            r = p.add_run(clean.upper())
            r.bold = True; r.font.size = Pt(16); r.font.name = TNR
            title_written = True
            continue

        if m:
            num, body = m.group(1), m.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            r_num = p.add_run(f'{num}.  ')
            r_num.bold = True; r_num.font.size = Pt(12); r_num.font.name = TNR
            r_body = p.add_run(body)
            r_body.font.size = Pt(12); r_body.font.name = TNR
        elif clean.isupper() and len(clean) < 80:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(clean)
            r.bold = True; r.font.size = Pt(13); r.font.name = TNR
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(clean)
            r.font.size = Pt(12); r.font.name = TNR

    add_watermark(doc, RDXPER_WATERMARK_TEXT)

    os.makedirs('generated', exist_ok=True)
    safe = re.sub(r'[^\w\-]', '_', (doc_type or 'Legal_Draft')[:40]) or 'Legal_Draft'
    out  = os.path.abspath(f'generated/{safe}_{uuid.uuid4().hex[:8]}.docx')
    doc.save(out)
    return out


def extract_text_from_upload(file_storage) -> str:
    """Extract plain text from an uploaded .docx or .txt reference format file."""
    filename = (file_storage.filename or '').lower()
    if filename.endswith('.docx'):
        tmp_path = os.path.abspath(f'generated/_upload_{uuid.uuid4().hex[:8]}.docx')
        os.makedirs('generated', exist_ok=True)
        file_storage.save(tmp_path)
        try:
            src = Document(tmp_path)
            text = '\n'.join(p.text for p in src.paragraphs if p.text.strip())
            for tbl in src.tables:
                for row in tbl.rows:
                    text += '\n' + ' | '.join(c.text for c in row.cells)
            return text
        finally:
            try: os.remove(tmp_path)
            except OSError: pass
    elif filename.endswith('.txt'):
        raw = file_storage.read()
        try:
            return raw.decode('utf-8')
        except UnicodeDecodeError:
            return raw.decode('latin-1', errors='ignore')
    else:
        raise ValueError('Unsupported file type — please upload a .docx or .txt file.')


def ai_draft_legal_document(doc_type: str, details: str, reference_text: str = '') -> str:
    """Call the AI model to draft a full legal document as plain text."""
    system = (
        'You are an expert legal drafter. Draft complete, professional, ready-to-use legal '
        'documents in plain text (no markdown, no asterisks, no code fences). '
        'Structure: a centred ALL-CAPS title on the first line, then the preamble/recitals '
        'as plain paragraphs, then the operative clauses as a numbered list ("1. ", "2. ", ...), '
        'and finally a signature block. Use precise, formal legal language appropriate for the '
        'jurisdiction implied by the details given. Do not include any commentary, explanations, '
        'or notes outside the document itself — output ONLY the document text.'
    )
    if reference_text:
        prompt = (
            f'Use the following document as the FORMAT/STRUCTURE reference — follow its layout, '
            f'clause structure and drafting style closely, but replace all names, dates, amounts '
            f'and other details with the DATA provided below. Fill in any gaps sensibly.\n\n'
            f'--- FORMAT REFERENCE ---\n{reference_text[:6000]}\n\n'
            f'--- DATA TO USE ---\n{details}\n\n'
            f'Now produce the complete final document text.'
        )
    else:
        prompt = (
            f'Draft a "{doc_type}" document using the following details and data:\n\n'
            f'{details}\n\n'
            f'Produce the complete, professional, ready-to-use document text.'
        )
    return ai_generate(prompt, system=system, temperature=0.4)


def build_trademark_license_docx(data: dict) -> str:
    """Generate a Licence to Use Trade Mark agreement as a .docx file
    (structured template — an alternative to the free-form AI drafting flow)."""
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    TNR = 'Times New Roman'

    def para(text, bold=False, sz=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sp_b=6, sp_a=6, center=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
        p.paragraph_format.space_before = Pt(sp_b)
        p.paragraph_format.space_after  = Pt(sp_a)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(sz)
        r.font.name = TNR
        return p

    def clause(number, text, sp_b=4, sp_a=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(sp_b)
        p.paragraph_format.space_after  = Pt(sp_a)
        p.paragraph_format.left_indent  = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        r_num = p.add_run(f'{number}.  ')
        r_num.bold = True
        r_num.font.size = Pt(12)
        r_num.font.name = TNR
        r_body = p.add_run(text)
        r_body.font.size = Pt(12)
        r_body.font.name = TNR
        return p

    para('LICENCE TO USE TRADE MARK', bold=True, sz=16, center=True, sp_b=0, sp_a=16)

    deed_date = data.get('deed_date', '').strip()
    if not deed_date:
        deed_date = datetime.now().strftime('%-d %B %Y')
    licensor_name = data.get('licensor_name', '[Licensor Name]')
    licensee_name = data.get('licensee_name', '[Licensee Name]')
    trademark     = data.get('trademark', '[TRADE MARK]')
    goods_services= data.get('goods_services', '[goods/services]')
    territory     = data.get('territory', '[Territory]')
    licence_fee_pct = data.get('licence_fee_pct', '10')
    payment_dates = '30th June and 31st December'
    notice_period = '3'

    preamble = (
        f'THIS DEED OF LICENCE is made on this {deed_date} between {licensor_name}, '
        f'hereinafter called the LICENSOR (which term shall unless excluded '
        f'by or repugnant to the context include its successors and assigns) of the one part '
        f'and {licensee_name}, '
        f'hereinafter referred to as the LICENSEE (which term '
        f'shall unless excluded by or repugnant to the context include its permitted nominees) '
        f'of the other part.'
    )
    para(preamble, sp_b=0, sp_a=10)

    recitals = [
        (f'WHEREAS the LICENSOR is the manufacturer of and dealer in {goods_services} and holds '
         f'the registered Trade Mark {trademark} in respect of {goods_services}.'),
        (f'AND WHEREAS the LICENSOR intends to expand its business and sell its products under its '
         f'Trade Mark in {territory}.'),
        (f'AND WHEREAS the LICENSEE has a manufacturing/trading unit to deal in {goods_services}.'),
        (f'AND WHEREAS the LICENSEE has approached the LICENSOR to grant licence to use the '
         f"LICENSOR's Trade Mark {trademark} for sale of the products/services of the LICENSEE."),
        (f'AND WHEREAS the LICENSOR has agreed to allow the LICENSEE to use its said Trade Mark '
         f'{trademark} to sell/provide the LICENSEE\'s {goods_services} on certain terms and conditions.'),
    ]
    for rec in recitals:
        para(rec, sp_b=4, sp_a=4)

    para('NOW THEREFORE THESE PRESENTS witnesseth and the parties hereby agree as follows:', bold=True, sp_b=10, sp_a=8)

    clauses = [
        (1, f'The LICENSOR hereby doth grant to the LICENSEE non-exclusive right to use the '
            f"LICENSOR's Trade Mark {trademark} in {territory} for sale/provision of its "
            f'{goods_services} under the Trade Name {trademark}.'),
        (2, f'The use of the Trade Mark by the LICENSEE shall be confined only to the items/services '
            f'that may be manufactured or provided by the LICENSEE at its own premises or through '
            f'its authorised channels. The LICENSEE shall pay half-yearly to the LICENSOR a licence '
            f'fee at the rate of {licence_fee_pct}% on the turnover of business of the LICENSEE, '
            f'such payment to be made by {payment_dates} every year.'),
        (3, f'The LICENSEE shall comply with the requirements and provisions of all laws, rules and '
            f'regulations in relation to the manufacture, sale or provision of {goods_services} '
            f'under the said Trade Mark of the LICENSOR.'),
        (4, f'The LICENSEE shall manufacture and sell/provide {goods_services} under the said Trade '
            f'Mark {trademark} in accordance with the specifications, make-up, brand and packing that '
            f'the LICENSOR may from time to time intimate to the LICENSEE.'),
        (5, f"The LICENSOR shall have access to the LICENSEE's manufacturing/service unit and to "
            f"inspect the LICENSEE's books of accounts and other records at all reasonable times on "
            f'giving prior notice.'),
        (6, f'The LICENSEE agrees, declares and covenants not to use the said Trade Mark or advertise '
            f'or publish in newspapers, journals, labels or any other documents or packages or do '
            f'anything having the effect of diluting the distinctiveness of the Trade Mark of the '
            f'LICENSOR. The LICENSEE shall give indications either visually or phonetically to the '
            f'purchasing public that the LICENSEE is using the Trade Mark {trademark} as the licensee '
            f'of the LICENSOR.'),
        (7, f'The LICENSEE undertakes to compensate the LICENSOR and keep the LICENSOR harmless from '
            f'and indemnified against all claims, proceedings, losses, costs and expenses for any '
            f'wilful or negligent conduct of the LICENSEE in relation to the use of the Trade Mark '
            f'of the LICENSOR.'),
        (8, f'The LICENSEE shall not acquire any right of registration of the Trade Mark by virtue '
            f'of the LICENSEE manufacturing, selling or providing {goods_services} as user of the '
            f'Trade Mark {trademark} for any number of years or after termination of the licence or otherwise.'),
        (9, f"The LICENSEE shall inform the LICENSOR of any infringement of the LICENSOR's Trade "
            f'Mark {trademark} with particulars of the infringement or passing off and the names and '
            f'addresses of the offenders.'),
        (10, f'The LICENSOR shall take and/or permit the LICENSEE to take all possible legal steps '
             f'for the protection and preservation of the Trade Mark and prevention of its '
             f'infringement or passing off by any person.'),
        (11, f'This agreement is terminable by giving {notice_period} months\' notice from either side.'),
        (12, f'In any legal proceedings or in any action against the infringement or passing off in '
             f'relation to the Trade Mark of the goods/services covered by the Licence, the LICENSEE '
             f'will take appropriate steps to protect the interests of the LICENSOR and allow the '
             f'LICENSOR to take any legal action or steps and to join the LICENSEE as a party therein.'),
    ]
    for num, text in clauses:
        clause(num, text)

    para('', sp_b=8, sp_a=0)
    para('THE SCHEDULE', bold=True, center=True, sp_b=8, sp_a=8)
    para('IN WITNESS WHEREOF the parties herein have executed these presents on the day, month and '
         'year first above-written.', sp_b=0, sp_a=16)

    para('Signed, sealed and delivered by', sp_b=0, sp_a=4)
    para(f'The authorised representative of {licensor_name} in the presence of:', sp_b=0, sp_a=12)

    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(8)
    sig_p.paragraph_format.space_after  = Pt(4)
    r1 = sig_p.add_run('1. ________________________')
    r1.font.name = TNR; r1.font.size = Pt(12)
    r1 = sig_p.add_run('\t\t\tSignature: ________________________')
    r1.font.name = TNR; r1.font.size = Pt(12)

    sig_p2 = doc.add_paragraph()
    sig_p2.paragraph_format.space_before = Pt(8)
    sig_p2.paragraph_format.space_after  = Pt(4)
    r2 = sig_p2.add_run('2. ________________________')
    r2.font.name = TNR; r2.font.size = Pt(12)
    r2b = sig_p2.add_run('\t\t\tDate: ________________________')
    r2b.font.name = TNR; r2b.font.size = Pt(12)

    add_watermark(doc, RDXPER_WATERMARK_TEXT)

    os.makedirs('generated', exist_ok=True)
    safe = re.sub(r'[^\w\-]', '_', f'TM_Licence_{licensor_name[:25]}')
    out  = os.path.abspath(f'generated/{safe}_{uuid.uuid4().hex[:8]}.docx')
    doc.save(out)
    return out


# ═══════════════════════════════════════════════════════════════════════════════
#  FRONTEND (single-page app)
# ═══════════════════════════════════════════════════════════════════════════════

HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>RDXper Legal — AI Legal Drafting</title>
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',Arial,sans-serif;background:#ffffff;color:#111111;min-height:100vh}
:root{--bg:#ffffff;--surface:#ffffff;--surface2:#f5f5f5;--border:#d0d0d0;--accent:#111111;--text:#111111;--muted:#666666;--dim:#999999;--error:#cc0000;--r:10px}
.wrap{max-width:760px;margin:0 auto;padding:0 20px}
header{padding:18px 0;display:flex;align-items:center;justify-content:space-between;border-bottom:2px solid #111}
.logo{display:flex;align-items:center;gap:10px}
.logo-mark{width:32px;height:32px;background:#111;border-radius:6px;display:flex;align-items:center;justify-content:center;font-weight:900;font-size:12px;color:#fff}
.logo-text{font-size:20px;font-weight:900;letter-spacing:-0.5px;color:#111}
.logo-text span{color:#111}
.user-chip{display:flex;align-items:center;gap:8px;background:#f5f5f5;border:1px solid #d0d0d0;border-radius:40px;padding:5px 12px;font-size:13px}
.nav-links{display:flex;gap:8px;align-items:center}
.nav-btn{background:none;border:1px solid #d0d0d0;color:#666;padding:5px 12px;border-radius:6px;cursor:pointer;font-size:12px;transition:all .2s}
.nav-btn:hover{border-color:#111;color:#111;background:#f5f5f5}
.nav-btn.danger{border-color:#cc0000;color:#cc0000}
.screen{display:none}.screen.active{display:block}
.hero{padding:56px 0 32px;text-align:center}
h1{font-size:clamp(26px,5vw,44px);font-weight:900;line-height:1.1;margin-bottom:16px;color:#111}
h1 em{color:#111;font-style:normal;border-bottom:3px solid #111}
.sub{font-size:15px;color:#666;max-width:520px;margin:0 auto 32px}
.card{background:#fff;border:1.5px solid #d0d0d0;border-radius:var(--r);padding:32px;max-width:420px;margin:0 auto;width:100%}
.ct{font-size:20px;font-weight:700;margin-bottom:6px;color:#111}
.cs{font-size:14px;color:#666;margin-bottom:20px}
.btn{width:100%;padding:13px 20px;border-radius:8px;border:none;font-size:15px;font-weight:700;cursor:pointer;transition:all .15s;display:flex;align-items:center;justify-content:center;gap:8px;margin-bottom:10px}
.btn:disabled{opacity:.4;cursor:not-allowed}
.btn-p{background:#111;color:#fff;border:2px solid #111}
.btn-p:hover:not(:disabled){background:#333;transform:translateY(-1px);box-shadow:0 4px 12px rgba(0,0,0,.2)}
.btn-dl{background:#111;color:#fff;border:2px solid #111;box-shadow:0 2px 8px rgba(0,0,0,.15)}
.btn-dl:hover:not(:disabled){background:#333;transform:translateY(-2px);box-shadow:0 6px 20px rgba(0,0,0,.25)}
.btn-s{background:#fff;color:#111;border:1.5px solid #d0d0d0}
.btn-s:hover:not(:disabled){border-color:#111;background:#f5f5f5}
.fg{margin-bottom:16px}.fg label{display:block;font-size:13px;color:#555;margin-bottom:6px;font-weight:600}
.fg input{width:100%;background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;color:#111;font-size:14px;outline:none;transition:border-color .2s}
.fg input:focus{border-color:#111;background:#fff}
textarea{width:100%;background:#f9f9f9;border:1.5px solid #d0d0d0;border-radius:8px;padding:10px 14px;color:#111;font-size:13px;outline:none;transition:border-color .2s;resize:vertical;font-family:'Segoe UI',Arial,sans-serif;line-height:1.6}
textarea:focus{border-color:#111;background:#fff}
textarea::placeholder{color:#bbb;font-size:12px}
.notif{display:none;padding:10px 14px;border-radius:8px;font-size:13px;margin-bottom:14px}
.notif.show{display:block}
.notif.error{background:#fff0f0;border:1.5px solid #cc0000;color:#cc0000}
.tabs{display:flex;gap:0;margin-bottom:20px;border-bottom:2px solid #d0d0d0}
.tab{padding:10px 18px;font-size:13px;cursor:pointer;border-radius:0;color:#666;border:none;background:none;transition:all .2s;border-bottom:2px solid transparent;margin-bottom:-2px;font-weight:600}
.tab.active{color:#111;border-bottom:2px solid #111;font-weight:700}
.spin{width:14px;height:14px;border:2px solid rgba(255,255,255,.4);border-top-color:#fff;border-radius:50%;animation:spin .7s linear infinite;display:inline-block}
@keyframes spin{to{transform:rotate(360deg)}}
footer{text-align:center;padding:32px 0;color:#999;font-size:12px;border-top:1.5px solid #d0d0d0;margin-top:40px}
@media(max-width:480px){
  .wrap{padding:0 12px}
  header{padding:12px 0;flex-wrap:wrap;gap:8px}
  h1{font-size:26px}
  .card{padding:20px 16px}
  .tabs{overflow-x:auto;white-space:nowrap;-webkit-overflow-scrolling:touch}
}
</style>
</head>
<body>
<div class="wrap">
<header>
  <div class="logo">
    <div class="logo-mark">rx</div>
    <div class="logo-text">RD<span>Xper</span> Legal</div>
  </div>
  <div class="nav-links" id="nav-auth" style="display:none">
    <div class="user-chip" id="user-chip">User</div>
    <button class="nav-btn danger" onclick="logout()">Sign out</button>
  </div>
</header>

<!-- LOGIN -->
<div class="screen active" id="s-home">
  <div class="hero">
    <h1>AI <em>Legal</em><br>Drafting</h1>
    <div class="sub">Describe the document you need, or upload a format to follow, and let AI draft it for you.</div>
  </div>
  <div class="card">
    <div class="ct">Sign in to continue</div>
    <div class="cs">Enter your name and email to get started</div>
    <div id="n-login" class="notif"></div>
    <div class="fg"><label>Name</label><input type="text" id="login-name" placeholder="Your name"></div>
    <div class="fg"><label>Email</label><input type="email" id="login-email" placeholder="you@example.com"></div>
    <button class="btn btn-p" id="btn-login" onclick="doLogin()">Sign in</button>
  </div>
</div>

<!-- LEGAL DRAFTING -->
<div class="screen" id="s-legal">
  <div style="padding-top:28px;max-width:700px;margin:0 auto">
    <div style="margin-bottom:16px">
      <div style="font-size:20px;font-weight:900;color:#111">⚖️ AI Legal Drafting</div>
    </div>
    <div class="cs" style="margin-bottom:20px">Describe the document you need and the details to include, and RDXper's AI will draft it for you — or upload a format/sample document and we'll follow its structure using your data.</div>

    <div class="tabs" id="legal-tabs">
      <button class="tab active" onclick="legalSwitchTab('custom',this)">✍️ Describe &amp; Generate</button>
      <button class="tab" onclick="legalSwitchTab('format',this)">📎 Use My Own Format</button>
    </div>

    <div id="n-legal" class="notif"></div>

    <div id="legal-form-custom">
      <div class="fg"><label>What kind of drafting do you need?</label>
        <input type="text" id="ld-doctype" placeholder="e.g. Rental Agreement, NDA, Employment Contract, Power of Attorney, Trademark Licence">
      </div>
      <div class="fg"><label>Details &amp; data for the draft</label>
        <textarea id="ld-details" rows="10" placeholder="Provide everything the document needs — party names & addresses, dates, amounts, terms, obligations, governing law, jurisdiction, special clauses, etc.&#10;&#10;Example: Landlord: Rohan Mehta, 12 MG Road, Pune. Tenant: Aisha Khan, 45 Park St, Pune. Property: 2BHK Flat No. 301, Green Meadows, Baner, Pune. Monthly rent: ₹28,000, payable by the 5th of every month. Security deposit: ₹1,00,000. Lease term: 11 months from 1 August 2026. Notice period: 1 month for termination by either party."></textarea>
      </div>
      <button class="btn btn-p" id="btn-legal-gen" onclick="generateLegalDoc()">⬇ Generate Draft (.docx)</button>
    </div>

    <div id="legal-form-format" style="display:none">
      <div class="fg"><label>Upload a format / sample document</label>
        <input type="file" id="ld-format-file" accept=".docx,.txt" style="width:100%;padding:10px;border:1.5px dashed #b0b0b0;border-radius:8px;background:#fafafa;font-size:13px">
        <div style="font-size:11px;color:#999;margin-top:4px">Accepted: .docx or .txt. We'll follow its structure and clauses.</div>
      </div>
      <div class="fg"><label>Data to fill into that format</label>
        <textarea id="ld-format-details" rows="9" placeholder="Provide the specific data that should replace the placeholders/details in the uploaded format, party names, dates, amounts, terms, etc."></textarea>
      </div>
      <button class="btn btn-p" id="btn-legal-gen-format" onclick="generateLegalDocFromFormat()">⬇ Generate Draft (.docx)</button>
    </div>

    <div id="legal-done" style="display:none;text-align:center;padding:32px 0">
      <div style="font-size:48px;margin-bottom:12px">✅</div>
      <div class="ct">Draft ready!</div>
      <div class="cs" id="legal-done-sub">Your document has been generated.</div>
      <button class="btn btn-dl" id="btn-legal-dl" onclick="downloadLegal()" style="max-width:360px;margin:16px auto 8px">⬇ Download Draft (.docx)</button>
      <button class="btn btn-s" onclick="resetLegal()" style="max-width:200px;margin:0 auto">Generate Another</button>
    </div>
  </div>
</div>

<footer>RDXper Legal — AI-powered legal drafting. Not a substitute for professional legal advice.</footer>
</div>

<script>
let token = localStorage.getItem('rdxper_legal_token') || '';
let userName = localStorage.getItem('rdxper_legal_name') || '';

function show(id){
  document.querySelectorAll('.screen').forEach(s=>s.classList.remove('active'));
  document.getElementById(id).classList.add('active');
}

function refreshNav(){
  if(token){
    document.getElementById('nav-auth').style.display='flex';
    document.getElementById('user-chip').textContent = userName || 'User';
    show('s-legal');
  } else {
    document.getElementById('nav-auth').style.display='none';
    show('s-home');
  }
}

async function doLogin(){
  const name = document.getElementById('login-name').value.trim();
  const email = document.getElementById('login-email').value.trim();
  const n = document.getElementById('n-login');
  n.classList.remove('show');
  if(!email || email.indexOf('@')===-1){
    n.className='notif error show'; n.textContent='Please enter a valid email.'; return;
  }
  const btn = document.getElementById('btn-login');
  btn.disabled = true; btn.innerHTML = '<span class="spin"></span> Signing in...';
  try{
    const r = await fetch('/api/auth/login',{
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({name, email})
    });
    const d = await r.json();
    if(!d.success){ n.className='notif error show'; n.textContent=d.message||'Login failed.'; return; }
    token = d.token; userName = d.name;
    localStorage.setItem('rdxper_legal_token', token);
    localStorage.setItem('rdxper_legal_name', userName);
    refreshNav();
  }catch(e){
    n.className='notif error show'; n.textContent='Connection error. Please try again.';
  }finally{
    btn.disabled=false; btn.innerHTML='Sign in';
  }
}

async function logout(){
  try{ await fetch('/api/auth/logout',{method:'POST',headers:{'Authorization':'Bearer '+token}}); }catch(e){}
  token=''; userName='';
  localStorage.removeItem('rdxper_legal_token');
  localStorage.removeItem('rdxper_legal_name');
  refreshNav();
}

// ── LEGAL DRAFTING ────────────────────────────────────────────────────────────
let legalJobId = '';
let legalDlName = 'RDXper_Legal_Draft.docx';

function legalSwitchTab(mode, btn){
  document.querySelectorAll('#legal-tabs .tab').forEach(t=>t.classList.remove('active'));
  btn.classList.add('active');
  document.getElementById('legal-form-custom').style.display = mode==='custom' ? 'block' : 'none';
  document.getElementById('legal-form-format').style.display = mode==='format' ? 'block' : 'none';
  document.getElementById('legal-done').style.display='none';
  const n=document.getElementById('n-legal'); n.classList.remove('show');
}

function resetLegal(){
  legalJobId='';
  document.getElementById('legal-done').style.display='none';
  document.getElementById('legal-form-custom').style.display='block';
  document.getElementById('legal-form-format').style.display='none';
  document.querySelectorAll('#legal-tabs .tab').forEach((t,i)=>t.classList.toggle('active', i===0));
  ['ld-doctype','ld-details','ld-format-details'].forEach(id=>{const el=document.getElementById(id);if(el)el.value='';});
  const fEl=document.getElementById('ld-format-file'); if(fEl) fEl.value='';
  document.getElementById('n-legal').classList.remove('show');
}

function legalShowError(msg){
  const n=document.getElementById('n-legal');
  n.className='notif error show'; n.textContent=msg;
}

async function generateLegalDoc(){
  const g = id => (document.getElementById(id)||{}).value.trim()||'';
  const doc_type = g('ld-doctype'), details = g('ld-details');
  if(!doc_type || !details){
    legalShowError('Please enter the type of document and the details/data for the draft.');
    return;
  }
  const btn=document.getElementById('btn-legal-gen');
  btn.disabled=true; btn.innerHTML='<span class="spin"></span> Drafting with AI...';
  document.getElementById('n-legal').classList.remove('show');
  try{
    const r = await fetch('/api/legal/generate',{
      method:'POST',
      headers:{'Content-Type':'application/json','Authorization':'Bearer '+token},
      body:JSON.stringify({mode:'custom', doc_type, details})
    });
    const d = await r.json();
    if(!d.success){ legalShowError(d.message||'Generation failed.'); return; }
    legalJobId = d.job_id;
    legalDlName = (doc_type.replace(/[^\\w\\-]+/g,'_')||'RDXper_Legal_Draft') + '.docx';
    document.getElementById('legal-done-sub').textContent = 'Your ' + doc_type + ' has been generated.';
    document.getElementById('legal-form-custom').style.display='none';
    document.getElementById('legal-form-format').style.display='none';
    document.getElementById('legal-done').style.display='block';
  }catch(e){
    legalShowError('Connection error. Please try again.');
  }finally{
    btn.disabled=false; btn.innerHTML='⬇ Generate Draft (.docx)';
  }
}

async function generateLegalDocFromFormat(){
  const details = (document.getElementById('ld-format-details')||{}).value.trim()||'';
  const fileEl = document.getElementById('ld-format-file');
  const file = fileEl && fileEl.files && fileEl.files[0];
  if(!file){ legalShowError('Please upload a format/sample document (.docx or .txt).'); return; }
  if(!details){ legalShowError('Please enter the data to fill into the uploaded format.'); return; }
  const btn=document.getElementById('btn-legal-gen-format');
  btn.disabled=true; btn.innerHTML='<span class="spin"></span> Drafting with AI...';
  document.getElementById('n-legal').classList.remove('show');
  try{
    const fd = new FormData();
    fd.append('mode','format');
    fd.append('details', details);
    fd.append('format_file', file);
    const r = await fetch('/api/legal/generate',{
      method:'POST',
      headers:{'Authorization':'Bearer '+token},
      body: fd
    });
    const d = await r.json();
    if(!d.success){ legalShowError(d.message||'Generation failed.'); return; }
    legalJobId = d.job_id;
    legalDlName = 'RDXper_Legal_Draft.docx';
    document.getElementById('legal-done-sub').textContent = 'Your document has been generated from the uploaded format.';
    document.getElementById('legal-form-custom').style.display='none';
    document.getElementById('legal-form-format').style.display='none';
    document.getElementById('legal-done').style.display='block';
  }catch(e){
    legalShowError('Connection error. Please try again.');
  }finally{
    btn.disabled=false; btn.innerHTML='⬇ Generate Draft (.docx)';
  }
}

async function downloadLegal(){
  const btn=document.getElementById('btn-legal-dl');
  btn.disabled=true; btn.innerHTML='<span class="spin"></span> Downloading...';
  try{
    const r = await fetch('/api/download/'+legalJobId,{headers:{'Authorization':'Bearer '+token}});
    if(!r.ok) throw new Error('failed');
    const blob=await r.blob(), url=URL.createObjectURL(blob), a=document.createElement('a');
    a.href=url; a.download=legalDlName; a.click(); URL.revokeObjectURL(url);
  }catch(e){ alert('Download failed. Please try again.'); }
  finally{ btn.disabled=false; btn.innerHTML='⬇ Download Draft (.docx)'; }
}

refreshNav();
</script>
</body>
</html>"""


# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return Response(HTML, mimetype='text/html')


@app.route('/api/auth/dev', methods=['POST'])
@app.route('/api/auth/login', methods=['POST'])
def simple_login():
    """Simple name + email login — works in all environments."""
    data  = request.json or {}
    email = data.get('email', '').strip().lower()
    name  = data.get('name', '').strip() or email.split('@')[0]
    if not email or '@' not in email:
        return jsonify({'success': False, 'message': 'Valid email required'}), 400
    user_id = 'u_' + email.replace('@', '_').replace('.', '_')
    with get_db() as db:
        user = db.execute('SELECT * FROM users WHERE email=?', (email,)).fetchone()
        if user:
            db.execute("UPDATE users SET name=?,last_login=datetime('now') WHERE email=?", (name, email))
            user_id = user['id']
        else:
            db.execute("INSERT INTO users (id,email,name,last_login) VALUES (?,?,?,datetime('now'))",
                       (user_id, email, name))
    tok = secrets.token_urlsafe(32)
    session_set(tok, email)
    sessions[tok]['user_id'] = user_id
    sessions[tok]['name'] = name
    return jsonify({'success': True, 'token': tok, 'email': email, 'name': name})


@app.route('/api/auth/logout', methods=['POST'])
def logout():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    session_delete(tok)
    return jsonify({'success': True})


@app.route('/api/legal/generate', methods=['POST'])
def gen_ai_legal_draft():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    sess = session_get(tok)
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if not os.environ.get('GROQ_API_KEY', '').strip():
        return jsonify({'success': False,
                        'message': 'GROQ_API_KEY not set. Get a free key at https://console.groq.com'}), 400

    is_multipart = request.content_type and 'multipart/form-data' in request.content_type
    mode = (request.form.get('mode') if is_multipart else (request.json or {}).get('mode')) or 'custom'

    try:
        if mode == 'format':
            details = (request.form.get('details') or '').strip()
            if not details:
                return jsonify({'success': False, 'message': 'Please provide the data to fill into the format.'}), 400
            file_storage = request.files.get('format_file')
            if not file_storage or not file_storage.filename:
                return jsonify({'success': False, 'message': 'Please upload a format/sample document.'}), 400
            reference_text = extract_text_from_upload(file_storage)
            doc_type = 'Legal Draft'
            ai_text = ai_draft_legal_document(doc_type, details, reference_text=reference_text)
        else:
            data = request.json or {}
            doc_type = (data.get('doc_type') or '').strip()
            details  = (data.get('details') or '').strip()
            if not doc_type or not details:
                return jsonify({'success': False, 'message': 'Please provide the document type and details.'}), 400
            ai_text = ai_draft_legal_document(doc_type, details)

        path = build_ai_legal_docx(doc_type, ai_text)
        jid  = uuid.uuid4().hex
        jobs[jid] = {'file_path': path, 'topic': doc_type, 'owner_email': sess['email']}
        return jsonify({'success': True, 'job_id': jid})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/legal/trademark-license', methods=['POST'])
def gen_trademark_license():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    sess = session_get(tok)
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.json or {}
    try:
        path = build_trademark_license_docx(data)
        jid  = uuid.uuid4().hex
        jobs[jid] = {'file_path': path, 'topic': 'Trademark License Agreement', 'owner_email': sess['email']}
        return jsonify({'success': True, 'job_id': jid})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/download/<jid>')
def download_draft(jid):
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    sess = session_get(tok)
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    job = jobs.get(jid)
    if not job:
        return jsonify({'success': False, 'message': 'Job not found'}), 404
    if job['owner_email'] != sess['email'] and not is_admin(sess):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403

    fp = job.get('file_path')
    if not fp or not os.path.exists(fp):
        return jsonify({'success': False, 'message': 'File not found on server'}), 404

    slug = re.sub(r'[^\w\-]', '_', (job.get('topic') or jid)[:40])
    return send_file(fp, as_attachment=True,
                     download_name=f'rdxper_legal_{slug}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    os.makedirs('generated', exist_ok=True)

    groq_key = os.environ.get('GROQ_API_KEY', '').strip()
    key_str = '\u2713 Groq \u2014 ready!' if groq_key else '\u2717 NOT SET \u2014 see below'
    print('\n' + '=' * 60)
    print('  rdxper-legal  \u2014  Standalone AI Legal Drafting App')
    print('  Powered by Groq (free tier)')
    print('  Open browser:  http://127.0.0.1:8081')
    print(f'  GROQ_API_KEY: {key_str}')
    print('=' * 60 + '\n')
    if not groq_key:
        print('  Get your free Groq API key at https://console.groq.com')
        print('  then:  export GROQ_API_KEY=your_key_here   (Mac/Linux)')
        print('         set GROQ_API_KEY=your_key_here       (Windows)\n')

    port = int(os.environ.get('PORT', 8081))
    app.run(host='0.0.0.0', port=port, debug=False, threaded=True)
