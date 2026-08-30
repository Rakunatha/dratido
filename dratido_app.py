"""
Dratido — "Draft Till Done"
────────────────────────────────────────────────
AI-assisted legal drafting workspace styled after the Indian court filing
tradition: case registers, cause lists, stamped pleadings — rebuilt for the
AI era.

Pipeline:
  1. Groq API (FREE)  → drafts the full legal document as plain text /
                         powers the AI Intelligence panel (research,
                         citations, clause suggestions, risk checks, etc.)
  2. python-docx      → formats it into a polished, watermarked .docx
  3. SQLite           → users, sessions, case files, drafts, versions,
                         comments — so the Dashboard and Case File System
                         are backed by real, persisted data.

AI Provider:
  Groq (free tier) — https://console.groq.com
  set GROQ_API_KEY=your_key_here

Usage:
  python dratido_app.py
"""

import os, uuid, time, secrets, re, json, sqlite3
import urllib.request, urllib.parse
from datetime import datetime
from flask import Flask, request, jsonify, send_file, Response
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import xml.sax.saxutils as _sax

app = Flask(__name__)
app.secret_key = secrets.token_hex(32)

sessions = {}   # token -> {email, user_id, name}
jobs     = {}   # job_id -> {file_path, topic, owner_email}

ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', '')

APP_NAME    = 'Dratido'
APP_TAGLINE = 'Draft Till Done'

# ── SQLite DB ──────────────────────────────────────────────────────────────
DB_PATH = os.environ.get('DB_PATH', 'dratido.db')

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as db:
        db.executescript("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY, email TEXT UNIQUE NOT NULL,
                name TEXT,
                created_at TEXT DEFAULT (datetime('now')),
                last_login TEXT
            );
            CREATE TABLE IF NOT EXISTS sessions (
                token TEXT PRIMARY KEY,
                email TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS case_files (
                id TEXT PRIMARY KEY,
                owner_email TEXT NOT NULL,
                case_no TEXT, court TEXT, bench TEXT, parties TEXT,
                advocate TEXT, filing_date TEXT, doc_type TEXT,
                status TEXT DEFAULT 'Drafting',
                timeline TEXT DEFAULT '[]',
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS drafts (
                id TEXT PRIMARY KEY,
                owner_email TEXT NOT NULL,
                case_id TEXT,
                title TEXT, doc_type TEXT,
                status TEXT DEFAULT 'Drafting',
                content TEXT DEFAULT '',
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS draft_versions (
                id TEXT PRIMARY KEY,
                draft_id TEXT NOT NULL,
                content TEXT, note TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS draft_comments (
                id TEXT PRIMARY KEY,
                draft_id TEXT NOT NULL,
                author TEXT, text TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS templates (
                id TEXT PRIMARY KEY,
                owner_email TEXT NOT NULL,
                title TEXT, doc_type TEXT, content TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            );
        """)

init_db()
os.makedirs('generated', exist_ok=True)


def session_set(token: str, email: str):
    sessions[token] = {'email': email}
    try:
        with get_db() as db:
            db.execute('INSERT OR REPLACE INTO sessions (token, email) VALUES (?, ?)', (token, email))
    except Exception as e:
        print(f'[session_set] DB error: {e}')


def session_get(token: str):
    if not token:
        return None
    if token in sessions:
        return sessions[token]
    try:
        with get_db() as db:
            row = db.execute('SELECT email FROM sessions WHERE token=?', (token,)).fetchone()
            if row:
                email = row['email']
                user = db.execute('SELECT id, name FROM users WHERE email=?', (email,)).fetchone()
                sessions[token] = {
                    'email': email,
                    'user_id': user['id'] if user else email,
                    'name': user['name'] if user else '',
                }
                return sessions[token]
    except Exception as e:
        print(f'[session_get] DB error: {e}')
    return None


def session_delete(token: str):
    sessions.pop(token, None)
    try:
        with get_db() as db:
            db.execute('DELETE FROM sessions WHERE token=?', (token,))
    except Exception as e:
        print(f'[session_delete] DB error: {e}')


def is_admin(sess: dict) -> bool:
    return bool(sess) and bool(ADMIN_EMAIL) and sess.get('email', '').strip().lower() == ADMIN_EMAIL.strip().lower()


def require_session():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    return session_get(tok)


def row2dict(row):
    return dict(row) if row else None


# ═══════════════════════════════════════════════════════════════════════════════
#  AI CLIENT  (Groq — fast free inference)
# ═══════════════════════════════════════════════════════════════════════════════

_GROQ_PREFERRED_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
]


def _get_groq_models(api_key, requests_module):
    headers = {"Authorization": f"Bearer {api_key}"}
    try:
        resp = requests_module.get(
            "https://api.groq.com/openai/v1/models",
            headers=headers,
            timeout=20,
        )
        if resp.status_code != 200:
            return [], f"HTTP {resp.status_code} from Groq /models: {resp.text[:300]}"
        data = resp.json()
        models = data.get("data", []) if isinstance(data, dict) else []
        ids = []
        for item in models:
            if not isinstance(item, dict):
                continue
            model_id = item.get("id")
            if model_id and item.get("active", True):
                ids.append(model_id)
        return ids, None
    except Exception as e:
        return [], f"Could not query Groq /models: {e}"


def _select_groq_models(api_key, requests_module):
    preferred_override = os.environ.get("GROQ_MODEL", "").strip()
    available, discovery_error = _get_groq_models(api_key, requests_module)
    available_set = set(available)

    if preferred_override:
        selected = [preferred_override]
        selected.extend(m for m in _GROQ_PREFERRED_MODELS
                        if m != preferred_override and m in available_set)
    elif available:
        selected = [m for m in _GROQ_PREFERRED_MODELS if m in available_set]
        excluded = ("whisper", "guard", "safeguard", "compound")
        selected.extend(
            m for m in available
            if m not in selected and not any(x in m.lower() for x in excluded)
        )
    else:
        selected = [preferred_override] if preferred_override else list(_GROQ_PREFERRED_MODELS)

    return list(dict.fromkeys(selected)), discovery_error


def ai_generate(prompt: str, system: str = "", temperature: float = 0.7) -> str:
    """Call Groq API with requests library + exponential backoff on 429."""
    import requests as _req

    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("GROQ_API_KEY not set. Get a free key at https://console.groq.com")

    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type":  "application/json",
    }

    _GROQ_MODELS, discovery_error = _select_groq_models(api_key, _req)
    print(f"[Groq] Models selected for this key/project: {_GROQ_MODELS}")
    if discovery_error:
        print(f"[Groq] Model discovery warning: {discovery_error}")
    if not _GROQ_MODELS:
        raise RuntimeError(
            "No active Groq text models are available to this API key/project. "
            "Check Groq Project > Settings > Limits/Model Permissions and API key."
        )

    last_error = None

    for model in _GROQ_MODELS:
        payload = {
            "model":       model,
            "messages":    messages,
            "temperature": temperature,
            "max_completion_tokens": 4096,
            "stream":      False,
        }

        for attempt in range(3):
            try:
                resp = _req.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=90,
                )
            except _req.exceptions.Timeout:
                last_error = f"Timeout on {model}"
                print(f"[Groq] Timeout on {model}, trying next...")
                break
            except _req.exceptions.RequestException as e:
                last_error = f"Request error on {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            status = resp.status_code

            if status == 429:
                wait = 2 ** (attempt + 2)
                last_error = f"429 rate-limited on {model} (attempt {attempt+1})"
                print(f"[Groq] 429 on {model}, waiting {wait}s...")
                time.sleep(wait)
                continue

            if status in (400, 402, 404, 503):
                body = resp.text[:300]
                last_error = f"HTTP {status} on {model}: {body}"
                print(f"[Groq] {status} on {model} (skipping): {body[:120]}")
                break

            if status != 200:
                last_error = f"HTTP {status} on {model}: {resp.text[:300]}"
                print(f"[Groq] Unexpected {status} on {model}: {resp.text[:120]}")
                break

            try:
                data = resp.json()
            except Exception as e:
                last_error = f"JSON parse error on {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            if "error" in data:
                err = data["error"]
                last_error = f"API error on {model}: {err}"
                print(f"[Groq] {last_error}")
                err_str = str(err).lower()
                if "rate" in err_str or "quota" in err_str or "limit" in err_str:
                    wait = 2 ** (attempt + 2)
                    print(f"[Groq] Quota error, waiting {wait}s...")
                    time.sleep(wait)
                    continue
                break

            try:
                text = (data["choices"][0]["message"]["content"] or "").strip()
            except (KeyError, IndexError, TypeError) as e:
                last_error = f"Unexpected shape from {model}: {e}"
                print(f"[Groq] {last_error}")
                break

            if not text:
                last_error = f"Empty content from {model}"
                print(f"[Groq] {last_error}")
                break

            print(f"[Groq] \u2713 {model} ({len(text)} chars)")
            return text

        time.sleep(1)

    raise RuntimeError(
        f"All accessible Groq models failed. Last error: {last_error}. "
        "The application queried Groq /models first, so this error now reflects "
        "models visible to your current API key/project. Check Groq Project "
        "model permissions and API key at https://console.groq.com"
    )


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX BUILDING
# ═══════════════════════════════════════════════════════════════════════════════

DRATIDO_WATERMARK_TEXT = 'Dratido - Draft Till Done'


def add_watermark(doc, text: str = DRATIDO_WATERMARK_TEXT):
    """Insert a diagonal, semi-transparent watermark into the header of every
    section (the classic Word VML watermark technique), plus a small text
    credit line in the footer as a reliable fallback for viewers that don't
    render VML shapes."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN

    safe_text = _sax.escape(text, {'"': '&quot;'})

    watermark_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<v:shapetype id="_x0000_t136" coordsize="1600,21600" o:spt="136" adj="10800" '
        'path="m@7,0l@8,0m@5,21600l@6,21600e">'
        '<v:formulas>'
        '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 #0"/>'
        '<v:f eqn="sum 0 0 #1"/><v:f eqn="prod #1 2 1"/><v:f eqn="sum 21600 0 #1"/>'
        '<v:f eqn="if #0 #3 0"/><v:f eqn="if #0 21600 #1"/><v:f eqn="if #3 21600 #2"/>'
        '<v:f eqn="if #3 #1 21600"/><v:f eqn="mid #4 #5"/><v:f eqn="mid #6 #7"/><v:f eqn="val #0"/>'
        '</v:formulas>'
        '<v:path textpathok="t" o:connecttype="custom" '
        'o:connectlocs="@9,0;@10,10800;@9,21600;@8,10800" o:connectangles="270,180,90,0"/>'
        '<v:textpath on="t" fitshape="t"/>'
        '</v:shapetype>'
        '<v:shape id="DratidoWatermark" o:spid="_x0000_s2049" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:520pt;height:110pt;'
        'rotation:315;z-index:-251654144;mso-position-horizontal:center;'
        'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
        'mso-position-vertical-relative:margin" o:allowincell="f" fillcolor="#8B1E2D" stroked="f">'
        '<v:fill opacity=".18"/>'
        f'<v:textpath style="font-family:\'Calibri\';font-size:1pt" string="{safe_text}"/>'
        '</v:shape>'
        '</w:pict>'
    )

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        h_para.text = ''
        h_para.alignment = _ALIGN.CENTER
        run = h_para.add_run()
        r_el = run._r
        pict = parse_xml(watermark_xml)
        r_el.append(pict)

        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        f_para.text = ''
        f_para.alignment = _ALIGN.CENTER
        f_run = f_para.add_run(text)
        f_run.font.size = Pt(8)
        f_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
        f_run.italic = True


def build_ai_legal_docx(doc_type: str, ai_text: str) -> str:
    """Convert the AI-drafted plain-text legal document into a formatted,
    watermarked .docx file resembling a formal Indian court filing."""
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    TNR = 'Times New Roman'
    lines = [ln.rstrip() for ln in ai_text.strip().split('\n')]

    numbered_re   = re.compile(r'^\s*(\d{1,3})[\.\)]\s+(.*)$')
    title_written = False

    for ln in lines:
        stripped = ln.strip()
        if not stripped:
            continue
        clean = stripped.strip('#').strip()
        clean = re.sub(r'^\*\*(.*)\*\*$', r'\1', clean).strip()
        clean = clean.lstrip('*').strip()
        if not clean:
            continue

        m = numbered_re.match(clean)
        if not title_written and not m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            r = p.add_run(clean.upper())
            r.bold = True; r.font.size = Pt(16); r.font.name = TNR
            title_written = True
            continue

        if m:
            num, body = m.group(1), m.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            r_num = p.add_run(f'{num}.  ')
            r_num.bold = True; r_num.font.size = Pt(12); r_num.font.name = TNR
            r_body = p.add_run(body)
            r_body.font.size = Pt(12); r_body.font.name = TNR
        elif clean.isupper() and len(clean) < 80:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(clean)
            r.bold = True; r.font.size = Pt(13); r.font.name = TNR
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(clean)
            r.font.size = Pt(12); r.font.name = TNR

    add_watermark(doc, DRATIDO_WATERMARK_TEXT)

    os.makedirs('generated', exist_ok=True)
    safe = re.sub(r'[^\w\-]', '_', (doc_type or 'Legal_Draft')[:40]) or 'Legal_Draft'
    out  = os.path.abspath(f'generated/{safe}_{uuid.uuid4().hex[:8]}.docx')
    doc.save(out)
    return out


def extract_text_from_upload(file_storage) -> str:
    """Extract plain text from an uploaded .docx or .txt reference format file."""
    filename = (file_storage.filename or '').lower()
    if filename.endswith('.docx'):
        tmp_path = os.path.abspath(f'generated/_upload_{uuid.uuid4().hex[:8]}.docx')
        os.makedirs('generated', exist_ok=True)
        file_storage.save(tmp_path)
        try:
            src = Document(tmp_path)
            text = '\n'.join(p.text for p in src.paragraphs if p.text.strip())
            for tbl in src.tables:
                for row in tbl.rows:
                    text += '\n' + ' | '.join(c.text for c in row.cells)
            return text
        finally:
            try: os.remove(tmp_path)
            except OSError: pass
    elif filename.endswith('.txt'):
        raw = file_storage.read()
        try:
            return raw.decode('utf-8')
        except UnicodeDecodeError:
            return raw.decode('latin-1', errors='ignore')
    else:
        raise ValueError('Unsupported file type — please upload a .docx or .txt file.')


def ai_draft_legal_document(doc_type: str, details: str, reference_text: str = '') -> str:
    """Call the AI model to draft a full legal document as plain text."""
    system = (
        'You are an expert Indian legal drafter trained in court filing conventions. Draft '
        'complete, professional, ready-to-use legal documents in plain text (no markdown, no '
        'asterisks, no code fences). '
        'Structure: a centred ALL-CAPS title on the first line (e.g. naming the court, case '
        'number placeholder, and document type where appropriate), then the cause-title / '
        'parties / preamble as plain paragraphs, then the operative clauses or averments as a '
        'numbered list ("1. ", "2. ", ...), then a prayer clause where applicable, and finally '
        'a verification and signature block. Use precise, formal legal language appropriate to '
        'the jurisdiction implied by the details given. Do not include any commentary, '
        'explanations, or notes outside the document itself — output ONLY the document text.'
    )
    if reference_text:
        prompt = (
            f'Use the following document as the FORMAT/STRUCTURE reference — follow its layout, '
            f'clause structure and drafting style closely, but replace all names, dates, amounts '
            f'and other details with the DATA provided below. Fill in any gaps sensibly.\n\n'
            f'--- FORMAT REFERENCE ---\n{reference_text[:6000]}\n\n'
            f'--- DATA TO USE ---\n{details}\n\n'
            f'Now produce the complete final document text.'
        )
    else:
        prompt = (
            f'Draft a "{doc_type}" document using the following details and data:\n\n'
            f'{details}\n\n'
            f'Produce the complete, professional, ready-to-use document text.'
        )
    return ai_generate(prompt, system=system, temperature=0.4)


# ═══════════════════════════════════════════════════════════════════════════════
#  AI INTELLIGENCE — assist actions for the drafting workspace
# ═══════════════════════════════════════════════════════════════════════════════

ASSIST_ACTIONS = {
    'improve_paragraph': {
        'label': 'Improve this paragraph',
        'system': 'You are a senior legal editor. Improve clarity, precision and formal tone '
                  'of the given text without changing its legal meaning. Return ONLY the '
                  'revised text, no commentary.',
    },
    'formal_language': {
        'label': 'Convert to formal legal language',
        'system': 'You are an expert legal drafter. Rewrite the given text into precise, '
                  'formal Indian legal drafting language (e.g. "hereinafter", "the said", '
                  '"whereas" where natural). Return ONLY the revised text, no commentary.',
    },
    'check_formatting': {
        'label': 'Check legal formatting',
        'system': 'You are a court registry scrutiny clerk. Review the given draft for '
                  'formatting and structural issues typical of Indian court filings (cause '
                  'title, numbering, verification, signature block, prayer clause, paragraph '
                  'numbering). Return a short bullet list of specific issues found and fixes. '
                  'If none, say so briefly.',
    },
    'find_inconsistencies': {
        'label': 'Find inconsistencies',
        'system': 'You are a meticulous legal proofreader. Identify factual or logical '
                  'inconsistencies, contradictions, undefined terms, or mismatched dates/names '
                  'within the given draft. Return a short bullet list. If none, say so briefly.',
    },
    'add_precedents': {
        'label': 'Add relevant precedents',
        'system': 'You are an Indian legal researcher. Based on the given draft or facts, '
                  'suggest 3-5 potentially relevant reported case-law precedents (Indian courts) '
                  'that could strengthen the arguments, with a one-line note on relevance for '
                  'each. Clearly state these are suggestions for the drafter to verify and cite '
                  'accurately, not confirmed citations.',
    },
    'legal_research': {
        'label': 'Legal research',
        'system': 'You are an Indian legal research assistant. Given the topic or facts, '
                  'summarise the relevant statutory provisions, doctrines and considerations a '
                  'drafter should account for. Be concise and organised under short headings.',
    },
    'citation_assist': {
        'label': 'Citation assistance',
        'system': 'You are a legal citation assistant. Given the text, suggest how to format '
                  'citations correctly (statute sections, case citations) in standard Indian '
                  'legal citation style, and flag any citation that looks incomplete or malformed.',
    },
    'clause_suggestions': {
        'label': 'Clause suggestions',
        'system': 'You are an expert legal drafter. Given the draft or facts, suggest '
                  'additional clauses commonly required in this type of document that appear '
                  'to be missing (e.g. jurisdiction, limitation, indemnity, verification). '
                  'Return a short list, each with a one-line draft of the clause.',
    },
    'risk_check': {
        'label': 'Risk / inconsistency detection',
        'system': 'You are a risk-review counsel. Review the given draft and flag clauses or '
                  'omissions that could pose legal or procedural risk. Return a short, '
                  'prioritised bullet list.',
    },
    'summarize': {
        'label': 'Document summarization',
        'system': 'You are a legal summarist. Summarise the given document in plain, precise '
                  'language in under 200 words, covering parties, relief sought and key facts.',
    },
    'explain_provision': {
        'label': 'Explain this provision',
        'system': 'You are a legal explainer for advocates and law students. Explain the given '
                  'clause or provision in plain English: what it means, why it matters, and any '
                  'common pitfalls. Be concise.',
    },
    'draft_from_facts': {
        'label': 'Draft from facts',
        'system': 'You are an expert legal drafter. Given raw facts, draft a properly '
                  'structured legal paragraph or clause suitable for inclusion in a formal '
                  'Indian court filing. Return ONLY the drafted text.',
    },
}


def ai_assist(action: str, text: str, extra: str = '') -> str:
    cfg = ASSIST_ACTIONS.get(action)
    if not cfg:
        raise ValueError('Unknown AI action')
    prompt = text.strip()
    if extra.strip():
        prompt += f'\n\n--- ADDITIONAL CONTEXT ---\n{extra.strip()}'
    if not prompt:
        raise ValueError('Please provide some text or facts for the AI to work with.')
    return ai_generate(prompt, system=cfg['system'], temperature=0.4)


# ═══════════════════════════════════════════════════════════════════════════════
#  FRONTEND (single-page app)
# ═══════════════════════════════════════════════════════════════════════════════

HTML_HEAD = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Dratido — Draft Till Done</title>
<link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' rx='10' fill='%236B1420'/%3E%3Ctext x='32' y='42' font-size='30' text-anchor='middle' fill='%23E9C97B' font-family='Georgia,serif'%3ED%3C/text%3E%3C/svg%3E">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700;800&family=Source+Serif+4:opsz,wght@8..60,500;8..60,600&family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
:root{
  --maroon:#6B1420;--maroon-dark:#4E0E17;--maroon-light:#8C2534;
  --ivory:#F8F3E7;--paper:#FCF9F1;--beige:#EDE3CC;--beige-dark:#DCCFA9;
  --charcoal:#26221D;--ink:#2E2A24;--muted:#726A5B;--dim:#9A917D;
  --gold:#B8924A;--gold-light:#E9C97B;--gold-dark:#8C6B2E;
  --line:#D8CBA6;--rule:#CBBB8C;
  --ok:#3E6B4A;--warn:#A2661C;--err:#9B2C2C;
  --r-sm:6px;--r-md:10px;--r-lg:16px;
  --shadow-1:0 1px 3px rgba(38,34,29,.08),0 1px 2px rgba(38,34,29,.06);
  --shadow-2:0 8px 28px rgba(38,34,29,.14);
  --font-head:'Playfair Display',Georgia,'Times New Roman',serif;
  --font-serif:'Source Serif 4',Georgia,serif;
  --font-ui:'Inter',-apple-system,'Segoe UI',sans-serif;
}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
html,body{height:100%}
body{
  font-family:var(--font-ui);background:var(--ivory);color:var(--ink);min-height:100vh;
  background-image:
    linear-gradient(var(--ivory),var(--ivory)),
    repeating-linear-gradient(0deg,rgba(107,20,32,.02) 0px,rgba(107,20,32,.02) 1px,transparent 1px,transparent 34px);
  background-blend-mode:normal;
}
a{color:inherit}
::selection{background:var(--gold-light);color:var(--maroon-dark)}
::-webkit-scrollbar{width:10px;height:10px}
::-webkit-scrollbar-thumb{background:var(--beige-dark);border-radius:10px}
::-webkit-scrollbar-track{background:transparent}

/* ── Seal / brand mark ─────────────────────────────────────────────── */
.seal{width:38px;height:38px;border-radius:50%;flex:0 0 auto;display:flex;align-items:center;justify-content:center;
  background:radial-gradient(circle at 35% 30%,var(--maroon-light),var(--maroon) 60%,var(--maroon-dark));
  border:2px solid var(--gold);box-shadow:inset 0 0 0 1px rgba(233,201,123,.35), var(--shadow-1);position:relative}
.seal::before{content:'';position:absolute;inset:4px;border:1px dashed rgba(233,201,123,.55);border-radius:50%}
.seal span{font-family:var(--font-head);color:var(--gold-light);font-weight:800;font-size:16px;line-height:1}
.brand{display:flex;align-items:center;gap:11px}
.brand-text{display:flex;flex-direction:column;line-height:1.05}
.brand-name{font-family:var(--font-head);font-weight:800;font-size:19px;letter-spacing:.3px;color:var(--maroon-dark)}
.brand-tag{font-size:10px;letter-spacing:2.2px;text-transform:uppercase;color:var(--gold-dark);font-weight:600;margin-top:2px}

/* ── Buttons ────────────────────────────────────────────────────────── */
.btn{display:inline-flex;align-items:center;justify-content:center;gap:8px;font-family:var(--font-ui);
  font-weight:600;font-size:13.5px;border-radius:var(--r-sm);padding:10px 18px;cursor:pointer;
  border:1.5px solid transparent;transition:all .15s ease;white-space:nowrap}
.btn:disabled{opacity:.5;cursor:not-allowed}
.btn-primary{background:var(--maroon);color:var(--ivory);border-color:var(--maroon)}
.btn-primary:hover:not(:disabled){background:var(--maroon-dark);box-shadow:var(--shadow-1)}
.btn-gold{background:var(--gold);color:var(--maroon-dark);border-color:var(--gold-dark)}
.btn-gold:hover:not(:disabled){background:var(--gold-light)}
.btn-ghost{background:transparent;color:var(--maroon-dark);border-color:var(--line)}
.btn-ghost:hover:not(:disabled){background:var(--beige);border-color:var(--rule)}
.btn-outline{background:var(--paper);color:var(--ink);border-color:var(--line)}
.btn-outline:hover:not(:disabled){border-color:var(--maroon);color:var(--maroon)}
.btn-danger{background:transparent;color:var(--err);border-color:#DEC2C2}
.btn-danger:hover:not(:disabled){background:#FBEEEE}
.btn-sm{padding:6px 12px;font-size:12px}
.btn-block{width:100%}
.spin{width:14px;height:14px;border:2px solid rgba(255,255,255,.35);border-top-color:currentColor;border-radius:50%;animation:spin .7s linear infinite;display:inline-block}
@keyframes spin{to{transform:rotate(360deg)}}

/* ── Auth screen ────────────────────────────────────────────────────── */
#screen-auth{min-height:100vh;display:flex;align-items:center;justify-content:center;padding:24px;
  background:
    radial-gradient(circle at 15% 10%,rgba(107,20,32,.06),transparent 40%),
    radial-gradient(circle at 85% 90%,rgba(184,146,74,.10),transparent 45%),
    var(--ivory);}
.auth-card{width:100%;max-width:420px;background:var(--paper);border:1px solid var(--line);border-radius:var(--r-lg);
  box-shadow:var(--shadow-2);padding:38px 34px 30px;position:relative;overflow:hidden}
.auth-card::before{content:'';position:absolute;top:0;left:0;right:0;height:6px;
  background:linear-gradient(90deg,var(--maroon),var(--gold),var(--maroon))}
.auth-brand{display:flex;flex-direction:column;align-items:center;text-align:center;margin-bottom:22px}
.auth-brand .seal{width:56px;height:56px;margin-bottom:12px}
.auth-brand .seal span{font-size:24px}
.auth-brand .brand-name{font-size:28px}
.auth-brand .brand-tag{font-size:11px;margin-top:4px}
.auth-sub{font-size:13.5px;color:var(--muted);text-align:center;margin-bottom:26px;line-height:1.55}
.auth-sub b{color:var(--maroon-dark)}
.fg{margin-bottom:15px}
.fg label{display:block;font-size:12.5px;font-weight:600;color:var(--muted);margin-bottom:6px;letter-spacing:.2px}
.fg input,.fg select{width:100%;background:var(--ivory);border:1.5px solid var(--line);border-radius:var(--r-sm);
  padding:11px 13px;font-size:14px;font-family:var(--font-ui);color:var(--ink);outline:none;transition:border-color .15s}
.fg input:focus,.fg select:focus{border-color:var(--maroon)}
.rule-divider{display:flex;align-items:center;gap:10px;margin:18px 0;color:var(--dim);font-size:11px;text-transform:uppercase;letter-spacing:1.5px}
.rule-divider::before,.rule-divider::after{content:'';flex:1;height:1px;background:var(--line)}
.auth-foot{text-align:center;font-size:12px;color:var(--dim);margin-top:18px}
.notif{display:none;padding:10px 13px;border-radius:var(--r-sm);font-size:13px;margin-bottom:14px;border:1.4px solid transparent}
.notif.show{display:block}
.notif.error{background:#FBEEEE;border-color:#DEC2C2;color:var(--err)}
.notif.ok{background:#EBF3ED;border-color:#C7DECC;color:var(--ok)}

/* ── App shell ──────────────────────────────────────────────────────── */
#app-shell{display:none;min-height:100vh}
#app-shell.active{display:flex}
.sidebar{width:236px;flex:0 0 auto;background:var(--maroon-dark);color:var(--ivory);display:flex;flex-direction:column;
  padding:20px 14px;position:sticky;top:0;height:100vh}
.sidebar .brand{padding:6px 8px 20px;border-bottom:1px solid rgba(233,201,123,.18);margin-bottom:14px}
.sidebar .brand-name{color:var(--ivory);font-size:18px}
.sidebar .brand-tag{color:var(--gold-light);opacity:.85}
.side-nav{display:flex;flex-direction:column;gap:3px}
.side-link{display:flex;align-items:center;gap:11px;padding:10px 12px;border-radius:8px;color:rgba(248,243,231,.78);
  font-size:13.5px;font-weight:600;cursor:pointer;transition:all .15s;border:1px solid transparent}
.side-link:hover{background:rgba(233,201,123,.09);color:var(--ivory)}
.side-link.active{background:rgba(233,201,123,.16);color:var(--gold-light);border-color:rgba(233,201,123,.3)}
.side-link .ic{width:18px;text-align:center;font-size:15px}
.side-section-label{font-size:10px;letter-spacing:1.6px;text-transform:uppercase;color:rgba(248,243,231,.4);
  margin:18px 10px 8px;font-weight:700}
.sidebar-foot{margin-top:auto;padding-top:14px;border-top:1px solid rgba(233,201,123,.18)}
.side-user{display:flex;align-items:center;gap:9px;padding:8px;border-radius:8px}
.side-avatar{width:30px;height:30px;border-radius:50%;background:var(--gold);color:var(--maroon-dark);
  display:flex;align-items:center;justify-content:center;font-weight:800;font-size:13px;flex:0 0 auto}
.side-user-name{font-size:12.5px;font-weight:600;color:var(--ivory);line-height:1.2}
.side-user-email{font-size:10.5px;color:rgba(248,243,231,.55)}
.side-logout{margin-top:10px;width:100%;background:none;border:1px solid rgba(233,201,123,.25);color:rgba(248,243,231,.7);
  border-radius:7px;padding:7px;font-size:11.5px;cursor:pointer;font-family:var(--font-ui);font-weight:600}
.side-logout:hover{border-color:var(--gold-light);color:var(--gold-light)}

.main{flex:1 1 auto;min-width:0;display:flex;flex-direction:column}
.topbar{display:flex;align-items:center;gap:16px;padding:14px 28px;background:var(--paper);
  border-bottom:1px solid var(--line);position:sticky;top:0;z-index:20}
.topbar-title{font-family:var(--font-head);font-size:20px;font-weight:700;color:var(--maroon-dark);flex:0 0 auto}
.search-box{flex:1 1 auto;max-width:460px;position:relative}
.search-box input{width:100%;padding:9px 14px 9px 34px;border-radius:20px;border:1.5px solid var(--line);
  background:var(--ivory);font-size:13px;font-family:var(--font-ui);outline:none}
.search-box input:focus{border-color:var(--maroon)}
.search-box .ic{position:absolute;left:12px;top:50%;transform:translateY(-50%);color:var(--dim);font-size:13px}
.topbar-actions{margin-left:auto;display:flex;gap:10px;align-items:center}
.view{padding:26px 28px 60px;flex:1}
.view{display:none}.view.active{display:block}

.section-head{display:flex;align-items:baseline;justify-content:space-between;margin-bottom:18px;flex-wrap:wrap;gap:10px}
.section-title{font-family:var(--font-head);font-size:24px;font-weight:700;color:var(--maroon-dark)}
.section-sub{font-size:13px;color:var(--muted);margin-top:3px}

/* ── Cards / stats ──────────────────────────────────────────────────── */
.grid-4{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;margin-bottom:26px}
.grid-2{display:grid;grid-template-columns:1.4fr 1fr;gap:20px}
.stat-card{background:var(--paper);border:1px solid var(--line);border-radius:var(--r-md);padding:18px 20px;
  box-shadow:var(--shadow-1);position:relative;overflow:hidden}
.stat-card::after{content:'';position:absolute;left:0;top:0;bottom:0;width:4px;background:var(--gold)}
.stat-num{font-family:var(--font-head);font-size:30px;font-weight:700;color:var(--maroon-dark);line-height:1}
.stat-label{font-size:12px;color:var(--muted);margin-top:6px;font-weight:600;text-transform:uppercase;letter-spacing:.5px}

.card{background:var(--paper);border:1px solid var(--line);border-radius:var(--r-md);box-shadow:var(--shadow-1)}
.card-head{display:flex;align-items:center;justify-content:space-between;padding:16px 20px;border-bottom:1px solid var(--line)}
.card-head h3{font-family:var(--font-head);font-size:16.5px;font-weight:700;color:var(--maroon-dark)}
.card-body{padding:6px 8px}
.card-body.pad{padding:18px 20px}

.file-row{display:flex;align-items:center;gap:14px;padding:13px 14px;border-radius:8px;cursor:pointer;transition:background .12s}
.file-row:hover{background:var(--beige)}
.file-tab{width:5px;align-self:stretch;border-radius:3px;flex:0 0 auto}
.file-icon{width:34px;height:34px;border-radius:7px;background:var(--beige);display:flex;align-items:center;justify-content:center;
  font-size:15px;color:var(--maroon-dark);flex:0 0 auto;border:1px solid var(--line)}
.file-main{flex:1;min-width:0}
.file-title{font-size:13.5px;font-weight:700;color:var(--ink);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.file-meta{font-size:11.5px;color:var(--muted);margin-top:2px}
.status-pill{font-size:10.5px;font-weight:700;padding:3px 10px;border-radius:20px;letter-spacing:.3px;text-transform:uppercase;flex:0 0 auto}
.status-drafting{background:#F3E6C9;color:#8C6B2E}
.status-review{background:#DCE6F2;color:#2E5A8C}
.status-final{background:#DCEBDF;color:#2E6B45}
.empty-note{padding:26px 20px;text-align:center;color:var(--dim);font-size:13px}

.quick-actions{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:26px}
.qa-btn{background:var(--paper);border:1.5px solid var(--line);border-radius:var(--r-md);padding:18px 16px;
  display:flex;flex-direction:column;gap:8px;cursor:pointer;transition:all .15s;text-align:left}
.qa-btn:hover{border-color:var(--maroon);box-shadow:var(--shadow-1);transform:translateY(-1px)}
.qa-btn .ic{width:34px;height:34px;border-radius:8px;background:var(--maroon);color:var(--gold-light);
  display:flex;align-items:center;justify-content:center;font-size:16px}
.qa-btn .qt{font-size:13.5px;font-weight:700;color:var(--ink)}
.qa-btn .qs{font-size:11.5px;color:var(--muted)}

/* ── Case / draft detail lists ─────────────────────────────────────── */
.case-list{display:flex;flex-direction:column;gap:10px}
.case-card{background:var(--paper);border:1px solid var(--line);border-left:4px solid var(--gold);border-radius:var(--r-md);
  padding:15px 18px;cursor:pointer;transition:all .15s;box-shadow:var(--shadow-1)}
.case-card:hover{box-shadow:var(--shadow-2);transform:translateY(-1px)}
.case-top{display:flex;justify-content:space-between;align-items:flex-start;gap:10px}
.case-no{font-family:var(--font-serif);font-weight:600;color:var(--maroon-dark);font-size:13.5px;letter-spacing:.3px}
.case-title{font-size:14.5px;font-weight:700;color:var(--ink);margin-top:3px}
.case-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:8px 18px;margin-top:10px;font-size:12px;color:var(--muted)}
.case-grid b{color:var(--ink);font-weight:600}

/* ── Drafting workspace ─────────────────────────────────────────────── */
.workspace{display:grid;grid-template-columns:1fr 340px;gap:20px;align-items:start}
.doc-toolbar{display:flex;align-items:center;gap:8px;padding:10px 14px;background:var(--paper);
  border:1px solid var(--line);border-radius:var(--r-md) var(--r-md) 0 0;flex-wrap:wrap}
.doc-toolbar .tbtn{background:none;border:1px solid transparent;color:var(--muted);padding:6px 10px;border-radius:6px;
  font-size:12px;cursor:pointer;font-weight:600;font-family:var(--font-ui)}
.doc-toolbar .tbtn:hover{background:var(--beige);color:var(--maroon-dark)}
.doc-toolbar .tbtn.active{background:var(--beige);color:var(--maroon-dark);border-color:var(--rule)}
.doc-toolbar .sep{width:1px;height:20px;background:var(--line);margin:0 4px}
.doc-status-input{border:none;background:transparent;font-family:var(--font-head);font-weight:700;font-size:15px;
  color:var(--maroon-dark);outline:none;flex:1;min-width:120px}
.paper-sheet{background:var(--paper);border:1px solid var(--line);border-top:none;padding:0;min-height:640px;
  display:flex;flex-direction:column;box-shadow:var(--shadow-1)}
.paper-inner{background:
    linear-gradient(#fff,#fff) padding-box,
    var(--paper);
  margin:22px auto;width:92%;max-width:680px;min-height:560px;padding:50px 56px 60px;
  box-shadow:0 0 0 1px var(--line), var(--shadow-2);position:relative;
  background:#FFFEFA;
  background-image:repeating-linear-gradient(#FFFEFA 0px,#FFFEFA 27px,var(--line) 28px);
  font-family:var(--font-serif);font-size:14.5px;line-height:28px;color:var(--charcoal);outline:none;
}
.paper-inner::before{content:'';position:absolute;left:44px;top:0;bottom:0;width:1px;background:rgba(107,20,32,.15)}
.paper-inner .doc-title{text-align:center;font-weight:700;font-family:var(--font-head);font-size:17px;
  letter-spacing:.4px;display:block;margin-bottom:6px;color:var(--maroon-dark)}
.paper-inner:empty::before{content:'Start typing your draft, or ask the AI assistant to draft it for you →';
  color:var(--dim);font-style:italic;position:static;background:none;width:auto;left:auto}
.stamp{position:absolute;right:34px;top:34px;width:96px;height:96px;border:2.5px solid var(--maroon);border-radius:50%;
  display:flex;align-items:center;justify-content:center;transform:rotate(-14deg);opacity:.55;color:var(--maroon);
  font-family:var(--font-head);font-weight:800;font-size:10.5px;text-align:center;letter-spacing:1px;line-height:1.3;pointer-events:none}
.doc-footer-bar{display:flex;justify-content:space-between;align-items:center;padding:10px 16px;background:var(--paper);
  border:1px solid var(--line);border-top:none;border-radius:0 0 var(--r-md) var(--r-md);font-size:11.5px;color:var(--dim)}

/* ── AI panel ───────────────────────────────────────────────────────── */
.ai-panel{background:var(--paper);border:1px solid var(--line);border-radius:var(--r-md);box-shadow:var(--shadow-1);
  display:flex;flex-direction:column;position:sticky;top:80px;max-height:calc(100vh - 100px)}
.ai-head{display:flex;align-items:center;gap:9px;padding:15px 16px;border-bottom:1px solid var(--line)}
.ai-head .dot{width:8px;height:8px;border-radius:50%;background:var(--ok)}
.ai-head h3{font-family:var(--font-head);font-size:15px;color:var(--maroon-dark);font-weight:700}
.ai-tabs{display:flex;border-bottom:1px solid var(--line)}
.ai-tab{flex:1;padding:9px 6px;text-align:center;font-size:11.5px;font-weight:700;color:var(--muted);cursor:pointer;
  border-bottom:2px solid transparent}
.ai-tab.active{color:var(--maroon-dark);border-color:var(--maroon)}
.ai-body{padding:14px 16px;overflow-y:auto;flex:1}
.prompt-grp-label{font-size:10.5px;text-transform:uppercase;letter-spacing:1px;color:var(--dim);font-weight:700;
  margin:14px 0 8px}
.prompt-grp-label:first-child{margin-top:0}
.prompt-chip{display:flex;align-items:center;gap:9px;width:100%;text-align:left;background:var(--ivory);
  border:1.3px solid var(--line);border-radius:8px;padding:9px 11px;font-size:12.5px;font-weight:600;color:var(--ink);
  cursor:pointer;margin-bottom:7px;transition:all .13s;font-family:var(--font-ui)}
.prompt-chip:hover{border-color:var(--maroon);background:#fff;transform:translateX(2px)}
.prompt-chip .ic{color:var(--maroon);font-size:13px;flex:0 0 auto}
.ai-response{background:var(--ivory);border:1px solid var(--line);border-radius:8px;padding:12px 13px;margin-bottom:10px;
  font-size:12.5px;line-height:1.65;color:var(--ink);white-space:pre-wrap}
.ai-response .arh{display:flex;justify-content:space-between;align-items:center;margin-bottom:7px}
.ai-response .arh b{font-size:11px;color:var(--maroon-dark);text-transform:uppercase;letter-spacing:.4px}
.ai-response .arh button{background:none;border:none;color:var(--muted);font-size:11px;cursor:pointer;font-weight:600}
.ai-foot{border-top:1px solid var(--line);padding:12px 14px}
.ai-foot textarea{width:100%;background:var(--ivory);border:1.5px solid var(--line);border-radius:8px;padding:9px 11px;
  font-size:12.5px;font-family:var(--font-ui);resize:vertical;min-height:52px;outline:none;margin-bottom:8px}
.ai-foot textarea:focus{border-color:var(--maroon)}

/* ── Side panels: history / comments / attachments ─────────────────── */
.hist-item{display:flex;justify-content:space-between;align-items:center;padding:9px 11px;border-bottom:1px solid var(--line);font-size:12px}
.hist-item:last-child{border-bottom:none}
.hist-item .hnote{font-weight:600;color:var(--ink)}
.hist-item .htime{color:var(--dim);font-size:11px}
.comment-item{padding:10px 12px;border-bottom:1px solid var(--line)}
.comment-item:last-child{border-bottom:none}
.comment-author{font-size:12px;font-weight:700;color:var(--maroon-dark)}
.comment-time{font-size:10.5px;color:var(--dim);margin-left:6px;font-weight:400}
.comment-text{font-size:12.5px;color:var(--ink);margin-top:3px;line-height:1.5}

/* ── Modal ──────────────────────────────────────────────────────────── */
.modal-overlay{display:none;position:fixed;inset:0;background:rgba(38,34,29,.55);z-index:100;
  align-items:center;justify-content:center;padding:20px}
.modal-overlay.show{display:flex}
.modal{background:var(--paper);border-radius:var(--r-lg);max-width:640px;width:100%;max-height:88vh;overflow-y:auto;
  box-shadow:var(--shadow-2);border:1px solid var(--line)}
.modal.wide{max-width:820px}
.modal-head{display:flex;justify-content:space-between;align-items:center;padding:18px 22px;border-bottom:1px solid var(--line);
  position:sticky;top:0;background:var(--paper);z-index:2}
.modal-head h2{font-family:var(--font-head);font-size:19px;color:var(--maroon-dark);font-weight:700}
.modal-close{background:none;border:none;font-size:20px;color:var(--dim);cursor:pointer;line-height:1}
.modal-body{padding:22px}
.modal-foot{display:flex;justify-content:flex-end;gap:10px;padding:16px 22px;border-top:1px solid var(--line)}
.tabs-row{display:flex;gap:0;border-bottom:2px solid var(--line);margin-bottom:18px}
.tab-btn{padding:9px 16px;font-size:13px;font-weight:700;color:var(--muted);background:none;border:none;cursor:pointer;
  border-bottom:2px solid transparent;margin-bottom:-2px;font-family:var(--font-ui)}
.tab-btn.active{color:var(--maroon-dark);border-color:var(--maroon)}
.form-row-2{display:grid;grid-template-columns:1fr 1fr;gap:14px}
textarea.plain{width:100%;background:var(--ivory);border:1.5px solid var(--line);border-radius:var(--r-sm);
  padding:11px 13px;font-size:13.5px;font-family:var(--font-serif);outline:none;resize:vertical;line-height:1.6}
textarea.plain:focus{border-color:var(--maroon)}
.doc-preview-sheet{background:#FFFEFA;border:1px solid var(--line);box-shadow:0 0 0 1px var(--line);padding:44px 46px;
  font-family:var(--font-serif);font-size:13.5px;line-height:1.85;color:var(--charcoal);white-space:pre-wrap;position:relative}
.template-pick{display:grid;grid-template-columns:repeat(2,1fr);gap:10px;margin-bottom:16px}
.tpick{border:1.5px solid var(--line);border-radius:8px;padding:12px 14px;cursor:pointer;text-align:left;background:var(--ivory);font-family:var(--font-ui)}
.tpick.active{border-color:var(--maroon);background:#fff;box-shadow:var(--shadow-1)}
.tpick .tt{font-size:13px;font-weight:700;color:var(--ink)}
.tpick .ts{font-size:11px;color:var(--muted);margin-top:2px}

/* toast */
.toast-wrap{position:fixed;bottom:22px;right:22px;z-index:200;display:flex;flex-direction:column;gap:8px}
.toast{background:var(--charcoal);color:var(--ivory);padding:11px 16px;border-radius:8px;font-size:12.5px;
  box-shadow:var(--shadow-2);border-left:3px solid var(--gold);max-width:320px}
.toast.err{border-color:var(--err)}

@media(max-width:980px){
  .workspace{grid-template-columns:1fr}
  .ai-panel{position:static;max-height:none}
  .grid-4,.quick-actions{grid-template-columns:repeat(2,1fr)}
  .grid-2{grid-template-columns:1fr}
  .case-grid{grid-template-columns:repeat(2,1fr)}
  .sidebar{position:fixed;left:-240px;z-index:50;transition:left .2s;box-shadow:var(--shadow-2)}
  .sidebar.open{left:0}
  .topbar-title{display:none}
}
@media(max-width:560px){
  .grid-4,.quick-actions{grid-template-columns:1fr}
  .form-row-2{grid-template-columns:1fr}
  .view{padding:18px 14px 50px}
  .topbar{padding:12px 14px}
}
</style>
</head>
"""

HTML_BODY = """
<body>

<!-- ══════════════════════ AUTH SCREEN ══════════════════════ -->
<div id="screen-auth">
  <div class="auth-card">
    <div class="auth-brand">
      <div class="seal"><span>D</span></div>
      <div class="brand-name">DRATIDO</div>
      <div class="brand-tag">Draft&nbsp;Till&nbsp;Done</div>
    </div>
    <p class="auth-sub">The <b>Indian court filing office</b>, rebuilt for the AI era.<br>Sign in to your drafting desk.</p>
    <div id="n-auth" class="notif error"></div>
    <div class="fg"><label>Full name</label><input id="in-name" placeholder="e.g. Adv. R. Krishnan" autocomplete="name"></div>
    <div class="fg"><label>Email address</label><input id="in-email" type="email" placeholder="you@chambers.in" autocomplete="email"></div>
    <button class="btn btn-primary btn-block" id="btn-login" onclick="doLogin()">Enter the Registry →</button>
    <div class="rule-divider">Est. for the modern Indian practice</div>
    <p class="auth-foot">No password needed for this workspace preview. Your drafts, cases and templates are saved to your account.</p>
  </div>
</div>

<!-- ══════════════════════ APP SHELL ══════════════════════ -->
<div id="app-shell">
  <aside class="sidebar" id="sidebar">
    <div class="brand">
      <div class="brand" style="gap:10px">
        <div class="seal"><span>D</span></div>
        <div class="brand-text">
          <div class="brand-name">DRATIDO</div>
          <div class="brand-tag">Draft Till Done</div>
        </div>
      </div>
    </div>
    <nav class="side-nav">
      <div class="side-link active" data-view="dashboard" onclick="goView('dashboard')"><span class="ic">⌂</span> Dashboard</div>
      <div class="side-link" data-view="workspace" onclick="goView('workspace')"><span class="ic">✎</span> Drafting Workspace</div>
      <div class="side-link" data-view="cases" onclick="goView('cases')"><span class="ic">🗂</span> Case File System</div>
      <div class="side-link" data-view="templates" onclick="goView('templates')"><span class="ic">▤</span> Saved Templates</div>
    </nav>
    <div class="side-section-label">Quick Draft</div>
    <div class="side-link" onclick="openNewDraftModal()"><span class="ic">+</span> New Draft</div>

    <div class="sidebar-foot">
      <div class="side-user">
        <div class="side-avatar" id="side-avatar">A</div>
        <div>
          <div class="side-user-name" id="side-name">—</div>
          <div class="side-user-email" id="side-email">—</div>
        </div>
      </div>
      <button class="side-logout" onclick="doLogout()">Sign out</button>
    </div>
  </aside>

  <div class="main">
    <div class="topbar">
      <button class="btn btn-ghost btn-sm" style="display:none" id="btn-menu" onclick="document.getElementById('sidebar').classList.toggle('open')">☰</button>
      <div class="topbar-title" id="topbar-title">Dashboard</div>
      <div class="search-box">
        <span class="ic">🔍</span>
        <input id="global-search" placeholder="Search by Case No., Party, Advocate, Court or Document Type…" oninput="onGlobalSearch(this.value)">
      </div>
      <div class="topbar-actions">
        <button class="btn btn-primary" onclick="openNewDraftModal()">+ New Draft</button>
      </div>
    </div>

    <!-- ── DASHBOARD ── -->
    <div class="view active" id="view-dashboard">
      <div class="section-head">
        <div><div class="section-title">Good day, <span id="greet-name">Counsel</span></div>
          <div class="section-sub">Your drafting desk, case register and cause list — in one place.</div></div>
      </div>

      <div class="quick-actions">
        <div class="qa-btn" onclick="openNewDraftModal('Petition')"><div class="ic">✎</div><div class="qt">Draft Petition</div><div class="qs">Start from facts</div></div>
        <div class="qa-btn" onclick="openNewDraftModal('Written Statement')"><div class="ic">🗎</div><div class="qt">Written Statement</div><div class="qs">Respond to a plaint</div></div>
        <div class="qa-btn" onclick="openNewDraftModal('Affidavit')"><div class="ic">🖋</div><div class="qt">Draft Affidavit</div><div class="qs">Sworn statement</div></div>
        <div class="qa-btn" onclick="openNewDraftModal('Legal Notice')"><div class="ic">✉</div><div class="qt">Legal Notice</div><div class="qs">Pre-litigation notice</div></div>
      </div>

      <div class="grid-4" id="stat-cards">
        <div class="stat-card"><div class="stat-num" id="stat-drafting">0</div><div class="stat-label">Drafting</div></div>
        <div class="stat-card"><div class="stat-num" id="stat-review">0</div><div class="stat-label">In Review</div></div>
        <div class="stat-card"><div class="stat-num" id="stat-final">0</div><div class="stat-label">Final</div></div>
        <div class="stat-card"><div class="stat-num" id="stat-cases">0</div><div class="stat-label">Active Cases</div></div>
      </div>

      <div class="grid-2">
        <div class="card">
          <div class="card-head"><h3>Recent Drafts</h3><button class="btn btn-ghost btn-sm" onclick="goView('workspace')">View all</button></div>
          <div class="card-body" id="recent-drafts-list"><div class="empty-note">Loading…</div></div>
        </div>
        <div class="card">
          <div class="card-head"><h3>Active Cases</h3><button class="btn btn-ghost btn-sm" onclick="goView('cases')">View all</button></div>
          <div class="card-body" id="active-cases-list"><div class="empty-note">Loading…</div></div>
        </div>
      </div>
    </div>

    <!-- ── WORKSPACE ── -->
    <div class="view" id="view-workspace">
      <div class="section-head">
        <div><div class="section-title">AI Drafting Workspace</div>
          <div class="section-sub" id="ws-sub">Select a draft, or start a new one, to begin.</div></div>
        <div style="display:flex;gap:10px">
          <select id="ws-draft-select" class="btn btn-outline" style="font-weight:600" onchange="loadDraft(this.value)"><option value="">— My Drafts —</option></select>
          <button class="btn btn-primary" onclick="openNewDraftModal()">+ New Draft</button>
        </div>
      </div>

      <div class="workspace" id="workspace-area" style="display:none">
        <div>
          <div class="doc-toolbar">
            <input class="doc-status-input" id="draft-title-input" placeholder="Untitled Draft" onchange="renameDraft()">
            <select id="draft-status-select" class="btn btn-outline btn-sm" onchange="setDraftStatus(this.value)">
              <option value="Drafting">Drafting</option><option value="Review">Review</option><option value="Final">Final</option>
            </select>
            <div class="sep"></div>
            <button class="tbtn" onclick="saveVersion()">💾 Save Version</button>
            <button class="tbtn" onclick="openHistoryModal()">🕘 Version History</button>
            <button class="tbtn" onclick="openCompareModal()">⇄ Compare Drafts</button>
            <button class="tbtn" onclick="openCommentsModal()">💬 Comments</button>
            <button class="tbtn" onclick="openPreviewModal()">📄 Preview</button>
          </div>
          <div class="paper-sheet">
            <div class="paper-inner" id="editor" contenteditable="true" spellcheck="true"></div>
          </div>
          <div class="doc-footer-bar">
            <span id="doc-wordcount">0 words</span>
            <span id="doc-savestate">All changes saved</span>
          </div>
        </div>

        <div class="ai-panel">
          <div class="ai-head"><span class="dot"></span><h3>AI Drafting Assistant</h3></div>
          <div class="ai-tabs">
            <div class="ai-tab active" data-tab="draft" onclick="switchAiTab('draft')">Draft</div>
            <div class="ai-tab" data-tab="intel" onclick="switchAiTab('intel')">Intelligence</div>
          </div>
          <div class="ai-body" id="ai-body-draft">
            <div class="prompt-grp-label">Draft a document</div>
            <button class="prompt-chip" onclick="quickDraftPrompt('Petition')"><span class="ic">✎</span> Draft Petition</button>
            <button class="prompt-chip" onclick="quickDraftPrompt('Written Statement')"><span class="ic">🗎</span> Draft Written Statement</button>
            <button class="prompt-chip" onclick="quickDraftPrompt('Affidavit')"><span class="ic">🖋</span> Draft Affidavit</button>
            <button class="prompt-chip" onclick="quickDraftPrompt('Legal Notice')"><span class="ic">✉</span> Draft Legal Notice</button>
            <div class="prompt-grp-label">Refine current draft</div>
            <button class="prompt-chip" onclick="runAssist('improve_paragraph')"><span class="ic">✦</span> Improve this paragraph</button>
            <button class="prompt-chip" onclick="runAssist('add_precedents')"><span class="ic">⚖</span> Add relevant precedents</button>
            <button class="prompt-chip" onclick="runAssist('check_formatting')"><span class="ic">≡</span> Check legal formatting</button>
            <button class="prompt-chip" onclick="runAssist('find_inconsistencies')"><span class="ic">⚠</span> Find inconsistencies</button>
            <div id="ai-responses-draft"></div>
          </div>
          <div class="ai-body" id="ai-body-intel" style="display:none">
            <div class="prompt-grp-label">AI Intelligence</div>
            <button class="prompt-chip" onclick="runAssist('legal_research')"><span class="ic">🔎</span> Legal research</button>
            <button class="prompt-chip" onclick="runAssist('citation_assist')"><span class="ic">§</span> Citation assistance</button>
            <button class="prompt-chip" onclick="runAssist('clause_suggestions')"><span class="ic">✚</span> Clause suggestions</button>
            <button class="prompt-chip" onclick="runAssist('risk_check')"><span class="ic">⚑</span> Risk / inconsistency detection</button>
            <button class="prompt-chip" onclick="runAssist('summarize')"><span class="ic">▤</span> Document summarization</button>
            <button class="prompt-chip" onclick="runAssist('explain_provision')"><span class="ic">?</span> Explain this provision</button>
            <button class="prompt-chip" onclick="runAssist('draft_from_facts')"><span class="ic">✎</span> Draft from facts</button>
            <button class="prompt-chip" onclick="runAssist('formal_language')"><span class="ic">✒</span> Convert to formal legal language</button>
            <div id="ai-responses-intel"></div>
          </div>
          <div class="ai-foot">
            <textarea id="ai-extra-input" placeholder="Optional: add facts, context or a specific instruction for the AI…"></textarea>
            <button class="btn btn-primary btn-block btn-sm" id="btn-ai-run" onclick="runSelectedOrCustom()">Run on selected text</button>
          </div>
        </div>
      </div>

      <div class="card" id="workspace-empty" style="padding:50px 20px;text-align:center">
        <div style="font-size:34px;margin-bottom:10px">🗎</div>
        <div style="font-weight:700;color:var(--maroon-dark);margin-bottom:6px">No draft open</div>
        <div style="color:var(--muted);font-size:13px;margin-bottom:16px">Choose a recent draft above, or start a fresh one.</div>
        <button class="btn btn-primary" onclick="openNewDraftModal()">+ New Draft</button>
      </div>
    </div>

    <!-- ── CASE FILES ── -->
    <div class="view" id="view-cases">
      <div class="section-head">
        <div><div class="section-title">Case File System</div>
          <div class="section-sub">Every matter, organised the way the registry always kept it — case number, bench, parties and timeline.</div></div>
        <button class="btn btn-primary" onclick="openNewCaseModal()">+ New Case File</button>
      </div>
      <div class="case-list" id="case-list"><div class="empty-note">Loading…</div></div>
    </div>

    <!-- ── TEMPLATES ── -->
    <div class="view" id="view-templates">
      <div class="section-head">
        <div><div class="section-title">Saved Templates</div>
          <div class="section-sub">Common Indian court filings, ready to fill in.</div></div>
      </div>
      <div class="grid-4" id="template-grid"></div>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: NEW DRAFT ══════════════════════ -->
<div class="modal-overlay" id="modal-newdraft">
  <div class="modal">
    <div class="modal-head"><h2>New Draft</h2><button class="modal-close" onclick="closeModal('modal-newdraft')">✕</button></div>
    <div class="modal-body">
      <div id="n-newdraft" class="notif error"></div>
      <div class="template-pick" id="nd-template-pick"></div>
      <div class="fg"><label>Document title</label><input id="nd-title" placeholder="e.g. Petition — Sharma vs. State"></div>
      <div class="fg"><label>Link to case file (optional)</label><select id="nd-case-select"><option value="">— No case file —</option></select></div>
      <div class="fg"><label>Facts / details for the AI to draft from</label>
        <textarea class="plain" id="nd-details" rows="6" placeholder="Describe the parties, facts, relief sought and any specific clauses to include…"></textarea></div>
    </div>
    <div class="modal-foot">
      <button class="btn btn-outline" onclick="closeModal('modal-newdraft')">Cancel</button>
      <button class="btn btn-ghost" onclick="createBlankDraft()">Start Blank</button>
      <button class="btn btn-primary" id="btn-nd-generate" onclick="createAndGenerateDraft()">✦ Draft with AI</button>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: NEW CASE ══════════════════════ -->
<div class="modal-overlay" id="modal-newcase">
  <div class="modal">
    <div class="modal-head"><h2 id="nc-heading">New Case File</h2><button class="modal-close" onclick="closeModal('modal-newcase')">✕</button></div>
    <div class="modal-body">
      <div id="n-newcase" class="notif error"></div>
      <div class="form-row-2">
        <div class="fg"><label>Case number</label><input id="nc-caseno" placeholder="e.g. C.S. No. 214/2026"></div>
        <div class="fg"><label>Court</label><input id="nc-court" placeholder="e.g. High Court of Madras"></div>
      </div>
      <div class="form-row-2">
        <div class="fg"><label>Bench</label><input id="nc-bench" placeholder="e.g. Division Bench"></div>
        <div class="fg"><label>Advocate</label><input id="nc-advocate" placeholder="e.g. Adv. R. Krishnan"></div>
      </div>
      <div class="fg"><label>Parties</label><input id="nc-parties" placeholder="e.g. Ravi Sharma vs. State of Tamil Nadu"></div>
      <div class="form-row-2">
        <div class="fg"><label>Filing date</label><input id="nc-filingdate" type="date"></div>
        <div class="fg"><label>Document type</label><input id="nc-doctype" placeholder="e.g. Civil Suit"></div>
      </div>
      <div class="fg"><label>Status</label>
        <select id="nc-status"><option>Drafting</option><option>Review</option><option>Final</option></select></div>
    </div>
    <div class="modal-foot">
      <button class="btn btn-outline" onclick="closeModal('modal-newcase')">Cancel</button>
      <button class="btn btn-primary" onclick="saveCase()">Save Case File</button>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: CASE DETAIL ══════════════════════ -->
<div class="modal-overlay" id="modal-casedetail">
  <div class="modal wide">
    <div class="modal-head"><h2 id="cd-title">Case File</h2><button class="modal-close" onclick="closeModal('modal-casedetail')">✕</button></div>
    <div class="modal-body">
      <div class="case-grid" id="cd-fields" style="grid-template-columns:repeat(2,1fr);margin-bottom:18px;font-size:13px"></div>
      <div class="tabs-row">
        <button class="tab-btn active" data-t="index" onclick="switchCaseTab('index')">Document Index</button>
        <button class="tab-btn" data-t="timeline" onclick="switchCaseTab('timeline')">Case Timeline</button>
        <button class="tab-btn" data-t="status" onclick="switchCaseTab('status')">Update Status</button>
      </div>
      <div id="cd-tab-index"></div>
      <div id="cd-tab-timeline" style="display:none">
        <div id="cd-timeline-list" style="margin-bottom:14px"></div>
        <div class="form-row-2">
          <input id="cd-tl-date" type="date">
          <input id="cd-tl-note" placeholder="Timeline note, e.g. Filed before Registry">
        </div>
        <button class="btn btn-outline btn-sm" style="margin-top:10px" onclick="addTimelineEntry()">+ Add Entry</button>
      </div>
      <div id="cd-tab-status" style="display:none">
        <div class="fg"><label>Case status</label>
          <select id="cd-status-select"><option>Drafting</option><option>Review</option><option>Final</option></select></div>
        <button class="btn btn-primary btn-sm" onclick="updateCaseStatus()">Update</button>
        <button class="btn btn-danger btn-sm" style="margin-left:10px" onclick="deleteCase()">Delete Case File</button>
      </div>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: PREVIEW ══════════════════════ -->
<div class="modal-overlay" id="modal-preview">
  <div class="modal wide">
    <div class="modal-head"><h2>Legal Document Preview</h2><button class="modal-close" onclick="closeModal('modal-preview')">✕</button></div>
    <div class="modal-body">
      <div class="doc-preview-sheet" id="preview-content"></div>
    </div>
    <div class="modal-foot">
      <button class="btn btn-outline" onclick="closeModal('modal-preview')">Close</button>
      <button class="btn btn-primary" id="btn-download-docx" onclick="downloadCurrentDocx()">⬇ Download as .docx</button>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: VERSION HISTORY ══════════════════════ -->
<div class="modal-overlay" id="modal-history">
  <div class="modal">
    <div class="modal-head"><h2>Version History</h2><button class="modal-close" onclick="closeModal('modal-history')">✕</button></div>
    <div class="modal-body">
      <div id="history-list"><div class="empty-note">No saved versions yet.</div></div>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: COMPARE DRAFTS ══════════════════════ -->
<div class="modal-overlay" id="modal-compare">
  <div class="modal wide">
    <div class="modal-head"><h2>Compare Versions</h2><button class="modal-close" onclick="closeModal('modal-compare')">✕</button></div>
    <div class="modal-body">
      <div class="form-row-2" style="margin-bottom:16px">
        <select id="cmp-a"></select>
        <select id="cmp-b"></select>
      </div>
      <button class="btn btn-primary btn-sm" onclick="runCompare()">Compare</button>
      <div id="cmp-result" style="margin-top:16px"></div>
    </div>
  </div>
</div>

<!-- ══════════════════════ MODAL: COMMENTS ══════════════════════ -->
<div class="modal-overlay" id="modal-comments">
  <div class="modal">
    <div class="modal-head"><h2>Comments</h2><button class="modal-close" onclick="closeModal('modal-comments')">✕</button></div>
    <div class="modal-body">
      <div id="comments-list" style="margin-bottom:14px"><div class="empty-note">No comments yet.</div></div>
      <textarea class="plain" id="new-comment" rows="3" placeholder="Add a remark for co-counsel or your future self…"></textarea>
      <button class="btn btn-primary btn-sm" style="margin-top:10px" onclick="postComment()">Post Comment</button>
    </div>
  </div>
</div>

<div class="toast-wrap" id="toast-wrap"></div>
"""

HTML_SCRIPT = """
<script>
let token = '';
let currentUser = {};
let cases = [];
let drafts = [];
let currentDraft = null;   // full draft object incl. versions/comments
let currentJobId = null;   // last generated docx job id
let saveTimer = null;
let pendingTemplate = '';
let editorDirty = false;

const TEMPLATES = [
  {doc_type:'Petition', label:'Draft Petition', ic:'✎', desc:'Civil / writ / criminal petition'},
  {doc_type:'Written Statement', label:'Written Statement', ic:'🗎', desc:'Reply to a plaint'},
  {doc_type:'Affidavit', label:'Affidavit', ic:'🖋', desc:'Sworn statement of facts'},
  {doc_type:'Legal Notice', label:'Legal Notice', ic:'✉', desc:'Pre-litigation notice'},
  {doc_type:'Bail Application', label:'Bail Application', ic:'⚖', desc:'Regular / anticipatory bail'},
  {doc_type:'Vakalatnama', label:'Vakalatnama', ic:'§', desc:'Advocate authorisation'},
  {doc_type:'Rejoinder', label:'Rejoinder', ic:'↩', desc:'Reply to written statement'},
  {doc_type:'Legal Opinion', label:'Legal Opinion', ic:'◆', desc:'Advisory memorandum'},
];

// ── Utilities ──────────────────────────────────────────────────────────
function toast(msg, isErr){
  const wrap = document.getElementById('toast-wrap');
  const t = document.createElement('div');
  t.className = 'toast' + (isErr ? ' err' : '');
  t.textContent = msg;
  wrap.appendChild(t);
  setTimeout(()=>t.remove(), 4200);
}
function fmtDate(s){
  if(!s) return '—';
  try{ const d = new Date(s.replace(' ','T')+'Z'); return d.toLocaleDateString('en-IN',{day:'2-digit',month:'short',year:'numeric'}); }
  catch(e){ return s; }
}
function initials(name){
  if(!name) return '?';
  const parts = name.trim().split(/\\s+/);
  return (parts[0][0] + (parts[1] ? parts[1][0] : '')).toUpperCase();
}
function statusClass(s){
  s = (s||'Drafting').toLowerCase();
  if(s==='review') return 'status-review';
  if(s==='final') return 'status-final';
  return 'status-drafting';
}
async function api(path, opts){
  opts = opts || {};
  opts.headers = Object.assign({'Authorization':'Bearer '+token}, opts.headers||{});
  if(opts.body && !(opts.body instanceof FormData)){
    opts.headers['Content-Type'] = 'application/json';
  }
  const r = await fetch(path, opts);
  let d;
  try{ d = await r.json(); }catch(e){ d = {success:false, message:'Server error'}; }
  return d;
}

// ── Auth ───────────────────────────────────────────────────────────────
function showAuthError(msg){
  const n = document.getElementById('n-auth');
  n.textContent = msg; n.classList.add('show');
}
async function doLogin(){
  const name = document.getElementById('in-name').value.trim();
  const email = document.getElementById('in-email').value.trim();
  document.getElementById('n-auth').classList.remove('show');
  if(!email || !email.includes('@')){ showAuthError('Please enter a valid email address.'); return; }
  const btn = document.getElementById('btn-login');
  btn.disabled = true; btn.innerHTML = '<span class="spin"></span> Entering…';
  try{
    const d = await api('/api/auth/login', {method:'POST', body: JSON.stringify({name, email})});
    if(!d.success){ showAuthError(d.message || 'Login failed.'); return; }
    token = d.token; currentUser = {name:d.name, email:d.email};
    localStorage.setItem('dratido_token', token);
    localStorage.setItem('dratido_user', JSON.stringify(currentUser));
    enterApp();
  }catch(e){ showAuthError('Connection error. Please try again.'); }
  finally{ btn.disabled = false; btn.innerHTML = 'Enter the Registry →'; }
}
function doLogout(){
  api('/api/auth/logout', {method:'POST'});
  localStorage.removeItem('dratido_token');
  localStorage.removeItem('dratido_user');
  location.reload();
}
function tryResumeSession(){
  const t = localStorage.getItem('dratido_token');
  const u = localStorage.getItem('dratido_user');
  if(t && u){ token = t; currentUser = JSON.parse(u); enterApp(); }
}
function enterApp(){
  document.getElementById('screen-auth').style.display = 'none';
  document.getElementById('app-shell').classList.add('active');
  document.getElementById('side-name').textContent = currentUser.name || currentUser.email;
  document.getElementById('side-email').textContent = currentUser.email;
  document.getElementById('side-avatar').textContent = initials(currentUser.name || currentUser.email);
  document.getElementById('greet-name').textContent = (currentUser.name || currentUser.email).split(' ')[0];
  loadTemplates();
  refreshAll();
}

// ── Navigation ─────────────────────────────────────────────────────────
function goView(view){
  document.querySelectorAll('.view').forEach(v=>v.classList.remove('active'));
  document.getElementById('view-'+view).classList.add('active');
  document.querySelectorAll('.side-link[data-view]').forEach(l=>l.classList.toggle('active', l.dataset.view===view));
  const titles = {dashboard:'Dashboard', workspace:'AI Drafting Workspace', cases:'Case File System', templates:'Saved Templates'};
  document.getElementById('topbar-title').textContent = titles[view] || '';
  document.getElementById('sidebar').classList.remove('open');
  if(view==='workspace') refreshDraftSelect();
  if(view==='cases') renderCaseList(cases);
}
function onGlobalSearch(q){
  q = q.trim().toLowerCase();
  if(!q){ renderRecentDrafts(drafts); renderCaseList(cases); return; }
  const hitDraft = d => (d.title+' '+d.doc_type).toLowerCase().includes(q);
  const hitCase = c => (c.case_no+' '+c.parties+' '+c.advocate+' '+c.court+' '+c.doc_type).toLowerCase().includes(q);
  renderRecentDrafts(drafts.filter(hitDraft));
  renderCaseList(cases.filter(hitCase));
}

// ── Data loading ───────────────────────────────────────────────────────
async function refreshAll(){
  await Promise.all([refreshCases(), refreshDrafts()]);
  renderStats();
  renderRecentDrafts(drafts);
  renderActiveCases(cases);
  renderCaseList(cases);
  refreshDraftSelect();
}
async function refreshCases(){
  const d = await api('/api/cases');
  if(d.success) cases = d.cases;
}
async function refreshDrafts(){
  const d = await api('/api/drafts');
  if(d.success) drafts = d.drafts;
}
function renderStats(){
  document.getElementById('stat-drafting').textContent = drafts.filter(d=>d.status==='Drafting').length;
  document.getElementById('stat-review').textContent = drafts.filter(d=>d.status==='Review').length;
  document.getElementById('stat-final').textContent = drafts.filter(d=>d.status==='Final').length;
  document.getElementById('stat-cases').textContent = cases.length;
}
function renderRecentDrafts(list){
  const el = document.getElementById('recent-drafts-list');
  if(!list.length){ el.innerHTML = '<div class="empty-note">No drafts yet. Start your first one above.</div>'; return; }
  el.innerHTML = list.slice(0,8).map(d => `
    <div class="file-row" onclick="goView('workspace'); setTimeout(()=>loadDraft('${d.id}'),50)">
      <div class="file-tab" style="background:var(--gold)"></div>
      <div class="file-icon">🗎</div>
      <div class="file-main">
        <div class="file-title">${escapeHtml(d.title||'Untitled Draft')}</div>
        <div class="file-meta">${escapeHtml(d.doc_type||'Document')} · updated ${fmtDate(d.updated_at)}</div>
      </div>
      <span class="status-pill ${statusClass(d.status)}">${d.status}</span>
    </div>`).join('');
}
function renderActiveCases(list){
  const el = document.getElementById('active-cases-list');
  if(!list.length){ el.innerHTML = '<div class="empty-note">No case files yet. Create one to get organised.</div>'; return; }
  el.innerHTML = list.slice(0,8).map(c => `
    <div class="file-row" onclick="openCaseDetail('${c.id}')">
      <div class="file-tab" style="background:var(--maroon)"></div>
      <div class="file-icon">🗂</div>
      <div class="file-main">
        <div class="file-title">${escapeHtml(c.parties||c.case_no||'Untitled Case')}</div>
        <div class="file-meta">${escapeHtml(c.case_no||'—')} · ${escapeHtml(c.court||'—')}</div>
      </div>
      <span class="status-pill ${statusClass(c.status)}">${c.status}</span>
    </div>`).join('');
}
function renderCaseList(list){
  const el = document.getElementById('case-list');
  if(!list.length){ el.innerHTML = '<div class="empty-note">No case files match. Create a new one to begin the register.</div>'; return; }
  el.innerHTML = list.map(c => `
    <div class="case-card" onclick="openCaseDetail('${c.id}')">
      <div class="case-top">
        <div><div class="case-no">${escapeHtml(c.case_no||'Case No. —')}</div>
          <div class="case-title">${escapeHtml(c.parties||'Untitled matter')}</div></div>
        <span class="status-pill ${statusClass(c.status)}">${c.status}</span>
      </div>
      <div class="case-grid">
        <div>Court<br><b>${escapeHtml(c.court||'—')}</b></div>
        <div>Bench<br><b>${escapeHtml(c.bench||'—')}</b></div>
        <div>Advocate<br><b>${escapeHtml(c.advocate||'—')}</b></div>
        <div>Filing Date<br><b>${escapeHtml(c.filing_date||'—')}</b></div>
      </div>
    </div>`).join('');
}
function escapeHtml(s){
  return (s==null?'':String(s)).replace(/[&<>"']/g, m => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[m]));
}

// ── Templates ──────────────────────────────────────────────────────────
function loadTemplates(){
  document.getElementById('template-grid').innerHTML = TEMPLATES.map(t => `
    <div class="qa-btn" onclick="openNewDraftModal('${t.doc_type}')">
      <div class="ic">${t.ic}</div><div class="qt">${t.label}</div><div class="qs">${t.desc}</div>
    </div>`).join('');
}

// ── Modal helpers ──────────────────────────────────────────────────────
function openModal(id){ document.getElementById(id).classList.add('show'); }
function closeModal(id){ document.getElementById(id).classList.remove('show'); }

// ── New Draft modal ────────────────────────────────────────────────────
function openNewDraftModal(preset){
  pendingTemplate = preset || '';
  document.getElementById('n-newdraft').classList.remove('show');
  document.getElementById('nd-title').value = preset ? (preset+' — Draft') : '';
  document.getElementById('nd-details').value = '';
  const pick = document.getElementById('nd-template-pick');
  pick.innerHTML = TEMPLATES.map(t => `
    <div class="tpick ${t.doc_type===preset?'active':''}" data-dt="${t.doc_type}" onclick="pickTemplate('${t.doc_type}')">
      <div class="tt">${t.ic} ${t.label}</div><div class="ts">${t.desc}</div>
    </div>`).join('');
  const sel = document.getElementById('nd-case-select');
  sel.innerHTML = '<option value="">— No case file —</option>' + cases.map(c=>`<option value="${c.id}">${escapeHtml(c.case_no||c.parties)}</option>`).join('');
  openModal('modal-newdraft');
}
function pickTemplate(dt){
  pendingTemplate = dt;
  document.querySelectorAll('.tpick').forEach(el=>el.classList.toggle('active', el.dataset.dt===dt));
  if(!document.getElementById('nd-title').value.trim()) document.getElementById('nd-title').value = dt + ' — Draft';
}
async function createBlankDraft(){
  const title = document.getElementById('nd-title').value.trim() || (pendingTemplate||'Untitled Draft');
  const case_id = document.getElementById('nd-case-select').value || null;
  const d = await api('/api/drafts', {method:'POST', body: JSON.stringify({title, doc_type: pendingTemplate, case_id})});
  if(!d.success){ toast(d.message||'Could not create draft.', true); return; }
  closeModal('modal-newdraft');
  await refreshDrafts(); refreshDraftSelect();
  goView('workspace');
  loadDraft(d.id);
}
async function createAndGenerateDraft(){
  const nEl = document.getElementById('n-newdraft'); nEl.classList.remove('show');
  const title = document.getElementById('nd-title').value.trim() || (pendingTemplate||'Untitled Draft');
  const details = document.getElementById('nd-details').value.trim();
  const case_id = document.getElementById('nd-case-select').value || null;
  if(!pendingTemplate){ nEl.textContent='Please choose a document type.'; nEl.classList.add('show'); return; }
  if(!details){ nEl.textContent='Please describe the facts and details for the AI to draft from.'; nEl.classList.add('show'); return; }
  const btn = document.getElementById('btn-nd-generate');
  btn.disabled = true; btn.innerHTML = '<span class="spin"></span> Drafting with AI…';
  try{
    const cd = await api('/api/drafts', {method:'POST', body: JSON.stringify({title, doc_type: pendingTemplate, case_id})});
    if(!cd.success) throw new Error(cd.message||'Could not create draft.');
    const gd = await api('/api/legal/generate', {method:'POST', body: JSON.stringify({doc_type: pendingTemplate, details, draft_id: cd.id})});
    if(!gd.success) throw new Error(gd.message||'Generation failed.');
    currentJobId = gd.job_id;
    closeModal('modal-newdraft');
    await refreshDrafts(); refreshDraftSelect();
    goView('workspace');
    await loadDraft(cd.id);
    toast('Draft generated by AI. Review and refine below.');
  }catch(e){ nEl.textContent = e.message || 'Connection error.'; nEl.classList.add('show'); }
  finally{ btn.disabled=false; btn.innerHTML='✦ Draft with AI'; }
}

// ── Workspace: draft loading & editing ────────────────────────────────
function refreshDraftSelect(){
  const sel = document.getElementById('ws-draft-select');
  const cur = sel.value;
  sel.innerHTML = '<option value="">— My Drafts —</option>' + drafts.map(d=>`<option value="${d.id}">${escapeHtml(d.title||'Untitled')} (${d.status})</option>`).join('');
  if(currentDraft) sel.value = currentDraft.id;
}
async function loadDraft(id){
  if(!id){ return; }
  const d = await api('/api/drafts/'+id);
  if(!d.success){ toast(d.message||'Could not load draft.', true); return; }
  currentDraft = d.draft;
  document.getElementById('workspace-area').style.display = 'grid';
  document.getElementById('workspace-empty').style.display = 'none';
  document.getElementById('ws-sub').textContent = 'Editing: ' + (currentDraft.title || 'Untitled Draft');
  document.getElementById('draft-title-input').value = currentDraft.title || '';
  document.getElementById('draft-status-select').value = currentDraft.status || 'Drafting';
  const editor = document.getElementById('editor');
  editor.innerHTML = renderEditorHtml(currentDraft.content || '');
  updateWordCount();
  document.getElementById('doc-savestate').textContent = 'All changes saved';
  refreshDraftSelect();
  document.getElementById('ai-responses-draft').innerHTML = '';
  document.getElementById('ai-responses-intel').innerHTML = '';
}
function renderEditorHtml(text){
  if(!text) return '';
  const lines = text.split('\\n');
  return lines.map((ln,i)=>{
    ln = escapeHtml(ln);
    if(i===0 && ln.trim()) return '<span class="doc-title">'+ln+'</span>';
    return ln || '<br>';
  }).join('<br>');
}
function editorPlainText(){
  const editor = document.getElementById('editor');
  return editor.innerText || '';
}
function updateWordCount(){
  const words = editorPlainText().trim().split(/\\s+/).filter(Boolean).length;
  document.getElementById('doc-wordcount').textContent = words + ' word' + (words===1?'':'s');
}
document.addEventListener('DOMContentLoaded', ()=>{
  const editor = document.getElementById('editor');
  if(editor){
    editor.addEventListener('input', ()=>{
      updateWordCount();
      document.getElementById('doc-savestate').textContent = 'Saving…';
      editorDirty = true;
      clearTimeout(saveTimer);
      saveTimer = setTimeout(autosaveDraft, 900);
    });
  }
});
async function autosaveDraft(){
  if(!currentDraft) return;
  const content = editorPlainText();
  const d = await api('/api/drafts/'+currentDraft.id, {method:'PUT', body: JSON.stringify({content})});
  if(d.success){ currentDraft.content = content; document.getElementById('doc-savestate').textContent = 'All changes saved'; editorDirty=false; }
}
async function renameDraft(){
  if(!currentDraft) return;
  const title = document.getElementById('draft-title-input').value.trim() || 'Untitled Draft';
  await api('/api/drafts/'+currentDraft.id, {method:'PUT', body: JSON.stringify({title})});
  currentDraft.title = title;
  await refreshDrafts(); refreshDraftSelect(); renderRecentDrafts(drafts);
}
async function setDraftStatus(status){
  if(!currentDraft) return;
  await api('/api/drafts/'+currentDraft.id, {method:'PUT', body: JSON.stringify({status})});
  currentDraft.status = status;
  await refreshDrafts(); renderStats(); renderRecentDrafts(drafts);
  toast('Draft marked as '+status+'.');
}
async function saveVersion(){
  if(!currentDraft) return;
  const content = editorPlainText();
  const note = prompt('Label this version (optional):', 'Manual save') || 'Manual save';
  const d = await api('/api/drafts/'+currentDraft.id, {method:'PUT', body: JSON.stringify({content, save_version:true, version_note:note})});
  if(d.success){ toast('Version saved.'); }
}

// ── Version History / Compare ─────────────────────────────────────────
async function openHistoryModal(){
  if(!currentDraft) return;
  const d = await api('/api/drafts/'+currentDraft.id);
  const list = document.getElementById('history-list');
  const versions = (d.draft && d.draft.versions) || [];
  if(!versions.length){ list.innerHTML = '<div class="empty-note">No saved versions yet. Use "Save Version" to snapshot your draft.</div>'; }
  else{
    list.innerHTML = versions.map(v => `
      <div class="hist-item">
        <div><div class="hnote">${escapeHtml(v.note||'Snapshot')}</div><div class="htime">${fmtDate(v.created_at)}</div></div>
        <button class="btn btn-outline btn-sm" onclick="restoreVersion('${v.id}')">Restore</button>
      </div>`).join('');
  }
  openModal('modal-history');
}
async function restoreVersion(vid){
  const d = await api('/api/drafts/'+currentDraft.id+'/versions/'+vid);
  if(!d.success){ toast(d.message||'Could not load version.', true); return; }
  document.getElementById('editor').innerHTML = renderEditorHtml(d.version.content||'');
  updateWordCount();
  document.getElementById('doc-savestate').textContent = 'Saving…';
  clearTimeout(saveTimer); saveTimer = setTimeout(autosaveDraft, 400);
  closeModal('modal-history');
  toast('Version restored into the editor.');
}
async function openCompareModal(){
  if(!currentDraft) return;
  const d = await api('/api/drafts/'+currentDraft.id);
  const versions = (d.draft && d.draft.versions) || [];
  const opts = ['<option value="__current__">Current draft</option>'].concat(
    versions.map(v=>`<option value="${v.id}">${escapeHtml(v.note||'Snapshot')} — ${fmtDate(v.created_at)}</option>`));
  document.getElementById('cmp-a').innerHTML = opts.join('');
  document.getElementById('cmp-b').innerHTML = opts.join('');
  document.getElementById('cmp-result').innerHTML = '';
  openModal('modal-compare');
}
async function getVersionText(val){
  if(val==='__current__') return editorPlainText();
  const d = await api('/api/drafts/'+currentDraft.id+'/versions/'+val);
  return d.success ? (d.version.content||'') : '';
}
function wordDiff(a, b){
  const A = a.split(/(\\s+)/), B = b.split(/(\\s+)/);
  const m = A.length, n = B.length;
  const dp = Array.from({length:m+1}, ()=>new Array(n+1).fill(0));
  for(let i=m-1;i>=0;i--) for(let j=n-1;j>=0;j--)
    dp[i][j] = A[i]===B[j] ? dp[i+1][j+1]+1 : Math.max(dp[i+1][j], dp[i][j+1]);
  let i=0, j=0, out=[];
  while(i<m && j<n){
    if(A[i]===B[j]){ out.push({t:A[i],k:'same'}); i++; j++; }
    else if(dp[i+1][j] >= dp[i][j+1]){ out.push({t:A[i],k:'del'}); i++; }
    else { out.push({t:B[j],k:'add'}); j++; }
  }
  while(i<m){ out.push({t:A[i],k:'del'}); i++; }
  while(j<n){ out.push({t:B[j],k:'add'}); j++; }
  return out;
}
async function runCompare(){
  const a = document.getElementById('cmp-a').value, b = document.getElementById('cmp-b').value;
  const [ta, tb] = await Promise.all([getVersionText(a), getVersionText(b)]);
  const diff = wordDiff(ta, tb);
  const html = diff.map(part => {
    const t = escapeHtml(part.t);
    if(part.k==='add') return '<span style="background:#DCEBDF;color:#2E6B45;text-decoration:none">'+t+'</span>';
    if(part.k==='del') return '<span style="background:#FBEEEE;color:#9B2C2C;text-decoration:line-through">'+t+'</span>';
    return t;
  }).join('');
  document.getElementById('cmp-result').innerHTML = '<div class="doc-preview-sheet" style="white-space:pre-wrap">'+html+'</div>';
}

// ── Comments ───────────────────────────────────────────────────────────
async function openCommentsModal(){
  if(!currentDraft) return;
  const d = await api('/api/drafts/'+currentDraft.id);
  const comments = (d.draft && d.draft.comments) || [];
  const list = document.getElementById('comments-list');
  list.innerHTML = comments.length ? comments.map(c => `
    <div class="comment-item"><span class="comment-author">${escapeHtml(c.author)}</span><span class="comment-time">${fmtDate(c.created_at)}</span>
      <div class="comment-text">${escapeHtml(c.text)}</div></div>`).join('') : '<div class="empty-note">No comments yet.</div>';
  document.getElementById('new-comment').value = '';
  openModal('modal-comments');
}
async function postComment(){
  const text = document.getElementById('new-comment').value.trim();
  if(!text || !currentDraft) return;
  const d = await api('/api/drafts/'+currentDraft.id+'/comments', {method:'POST', body: JSON.stringify({text})});
  if(d.success){ openCommentsModal(); } else { toast(d.message||'Could not post comment.', true); }
}

// ── AI panel ───────────────────────────────────────────────────────────
function switchAiTab(tab){
  document.querySelectorAll('.ai-tab').forEach(t=>t.classList.toggle('active', t.dataset.tab===tab));
  document.getElementById('ai-body-draft').style.display = tab==='draft' ? 'block' : 'none';
  document.getElementById('ai-body-intel').style.display = tab==='intel' ? 'block' : 'none';
}
function getSelectedOrFullText(){
  const sel = window.getSelection();
  if(sel && sel.toString().trim().length > 0 && document.getElementById('editor').contains(sel.anchorNode)){
    return sel.toString();
  }
  return editorPlainText();
}
async function quickDraftPrompt(docType){
  if(!currentDraft){ openNewDraftModal(docType); return; }
  const details = prompt('Describe the facts/details for the AI to draft a "'+docType+'":', '');
  if(!details) return;
  await runGenerateInto(docType, details);
}
async function runGenerateInto(docType, details){
  const respWrap = document.getElementById('ai-responses-draft');
  const holder = document.createElement('div');
  holder.className = 'ai-response';
  holder.innerHTML = '<div class="arh"><b>Drafting…</b></div><span class="spin"></span>';
  respWrap.prepend(holder);
  try{
    const gd = await api('/api/legal/generate', {method:'POST', body: JSON.stringify({doc_type:docType, details, draft_id: currentDraft.id})});
    if(!gd.success) throw new Error(gd.message||'Generation failed.');
    currentJobId = gd.job_id;
    document.getElementById('editor').innerHTML = renderEditorHtml(gd.preview);
    updateWordCount();
    currentDraft.content = gd.preview;
    holder.innerHTML = '<div class="arh"><b>'+escapeHtml(docType)+' drafted</b></div>Inserted into the editor. Review, edit, then Save Version.';
    await refreshDrafts(); renderRecentDrafts(drafts);
  }catch(e){ holder.innerHTML = '<div class="arh"><b>Error</b></div>'+escapeHtml(e.message||'Something went wrong.'); }
}
async function runAssist(action){
  if(!currentDraft){ toast('Open or start a draft first.', true); return; }
  const text = getSelectedOrFullText();
  const extra = document.getElementById('ai-extra-input').value.trim();
  const tab = document.getElementById('ai-body-intel').style.display==='none' ? 'draft' : 'intel';
  const wrap = document.getElementById(tab==='draft' ? 'ai-responses-draft' : 'ai-responses-intel');
  const labelMap = {improve_paragraph:'Improve this paragraph', formal_language:'Convert to formal legal language',
    check_formatting:'Legal formatting check', find_inconsistencies:'Inconsistency check', add_precedents:'Relevant precedents',
    legal_research:'Legal research', citation_assist:'Citation assistance', clause_suggestions:'Clause suggestions',
    risk_check:'Risk detection', summarize:'Summary', explain_provision:'Explanation', draft_from_facts:'Drafted clause'};
  const holder = document.createElement('div');
  holder.className = 'ai-response';
  holder.innerHTML = '<div class="arh"><b>'+ (labelMap[action]||action) +'</b></div><span class="spin"></span> Thinking…';
  wrap.prepend(holder);
  try{
    const d = await api('/api/legal/assist', {method:'POST', body: JSON.stringify({action, text, extra})});
    if(!d.success) throw new Error(d.message||'AI request failed.');
    holder.innerHTML = '<div class="arh"><b>'+(labelMap[action]||action)+'</b><button onclick="insertIntoEditor(this)">Insert →</button></div>'
      + '<div class="assist-out" style="display:none">'+escapeHtml(d.result)+'</div>'+escapeHtml(d.result);
  }catch(e){ holder.innerHTML = '<div class="arh"><b>Error</b></div>'+escapeHtml(e.message||'Something went wrong.'); }
}
function insertIntoEditor(btn){
  const holder = btn.closest('.ai-response');
  const raw = holder.querySelector('.assist-out').textContent;
  const editor = document.getElementById('editor');
  editor.innerHTML += '<br>' + renderEditorHtml(raw);
  updateWordCount();
  clearTimeout(saveTimer); saveTimer = setTimeout(autosaveDraft, 400);
  toast('Inserted into draft.');
}
function runSelectedOrCustom(){
  const extra = document.getElementById('ai-extra-input').value.trim();
  if(!extra){ toast('Type an instruction, or use a prompt above.', true); return; }
  runAssist('draft_from_facts');
}

// ── Preview & download ────────────────────────────────────────────────
function openPreviewModal(){
  if(!currentDraft){ toast('Open a draft first.', true); return; }
  document.getElementById('preview-content').textContent = editorPlainText() || '(This draft is empty.)';
  openModal('modal-preview');
}
async function downloadCurrentDocx(){
  if(!currentDraft) return;
  const btn = document.getElementById('btn-download-docx');
  btn.disabled = true; btn.innerHTML = '<span class="spin"></span> Preparing…';
  try{
    const content = editorPlainText();
    const gd = await api('/api/legal/generate', {method:'POST', body: JSON.stringify({doc_type: currentDraft.doc_type||currentDraft.title, details: content || currentDraft.title, draft_id: currentDraft.id})});
    if(!gd.success) throw new Error(gd.message||'Could not prepare file.');
    currentJobId = gd.job_id;
    const r = await fetch('/api/download/'+currentJobId, {headers:{'Authorization':'Bearer '+token}});
    if(!r.ok) throw new Error('Download failed.');
    const blob = await r.blob(), url = URL.createObjectURL(blob), a = document.createElement('a');
    a.href = url; a.download = 'Dratido_'+(currentDraft.title||'Draft').replace(/[^\\w\\-]/g,'_')+'.docx'; a.click();
    URL.revokeObjectURL(url);
  }catch(e){ toast(e.message||'Download failed.', true); }
  finally{ btn.disabled=false; btn.innerHTML='⬇ Download as .docx'; }
}

// ── Case files ─────────────────────────────────────────────────────────
let editingCaseId = null;
function openNewCaseModal(){
  editingCaseId = null;
  document.getElementById('nc-heading').textContent = 'New Case File';
  document.getElementById('n-newcase').classList.remove('show');
  ['nc-caseno','nc-court','nc-bench','nc-advocate','nc-parties','nc-filingdate','nc-doctype'].forEach(id=>document.getElementById(id).value='');
  document.getElementById('nc-status').value = 'Drafting';
  openModal('modal-newcase');
}
async function saveCase(){
  const payload = {
    case_no: document.getElementById('nc-caseno').value.trim(),
    court: document.getElementById('nc-court').value.trim(),
    bench: document.getElementById('nc-bench').value.trim(),
    advocate: document.getElementById('nc-advocate').value.trim(),
    parties: document.getElementById('nc-parties').value.trim(),
    filing_date: document.getElementById('nc-filingdate').value,
    doc_type: document.getElementById('nc-doctype').value.trim(),
    status: document.getElementById('nc-status').value,
  };
  const n = document.getElementById('n-newcase');
  if(!payload.case_no && !payload.parties){ n.textContent='Please enter at least a case number or party names.'; n.classList.add('show'); return; }
  const d = await api('/api/cases', {method:'POST', body: JSON.stringify(payload)});
  if(!d.success){ n.textContent = d.message||'Could not save.'; n.classList.add('show'); return; }
  closeModal('modal-newcase');
  await refreshCases(); renderStats(); renderActiveCases(cases); renderCaseList(cases);
  toast('Case file created.');
}
let currentCaseId = null;
async function openCaseDetail(id){
  const d = await api('/api/cases/'+id);
  if(!d.success){ toast(d.message||'Could not load case.', true); return; }
  const c = d.case;
  currentCaseId = id;
  document.getElementById('cd-title').textContent = c.parties || c.case_no || 'Case File';
  document.getElementById('cd-fields').innerHTML = `
    <div>Case No.<br><b>${escapeHtml(c.case_no||'—')}</b></div>
    <div>Court<br><b>${escapeHtml(c.court||'—')}</b></div>
    <div>Bench<br><b>${escapeHtml(c.bench||'—')}</b></div>
    <div>Advocate<br><b>${escapeHtml(c.advocate||'—')}</b></div>
    <div>Filing Date<br><b>${escapeHtml(c.filing_date||'—')}</b></div>
    <div>Status<br><b>${escapeHtml(c.status||'—')}</b></div>`;
  const idx = document.getElementById('cd-tab-index');
  idx.innerHTML = (c.drafts||[]).length ? c.drafts.map(dr => `
    <div class="file-row" onclick="closeModal('modal-casedetail'); goView('workspace'); setTimeout(()=>loadDraft('${dr.id}'),50)">
      <div class="file-tab" style="background:var(--gold)"></div><div class="file-icon">🗎</div>
      <div class="file-main"><div class="file-title">${escapeHtml(dr.title)}</div><div class="file-meta">${escapeHtml(dr.doc_type||'')} · ${fmtDate(dr.updated_at)}</div></div>
      <span class="status-pill ${statusClass(dr.status)}">${dr.status}</span>
    </div>`).join('') : '<div class="empty-note">No documents linked yet. Start a draft and link it to this case.</div>';
  renderTimeline(c.timeline||[]);
  document.getElementById('cd-status-select').value = c.status || 'Drafting';
  switchCaseTab('index');
  openModal('modal-casedetail');
}
function renderTimeline(tl){
  const el = document.getElementById('cd-timeline-list');
  el.innerHTML = tl.length ? tl.map(t=>`<div class="hist-item"><div><div class="hnote">${escapeHtml(t.note)}</div><div class="htime">${escapeHtml(t.date)}</div></div></div>`).join('')
    : '<div class="empty-note">No timeline entries yet.</div>';
}
function switchCaseTab(t){
  document.querySelectorAll('#modal-casedetail .tab-btn').forEach(b=>b.classList.toggle('active', b.dataset.t===t));
  ['index','timeline','status'].forEach(k=>document.getElementById('cd-tab-'+k).style.display = k===t?'block':'none');
}
async function addTimelineEntry(){
  const date = document.getElementById('cd-tl-date').value;
  const note = document.getElementById('cd-tl-note').value.trim();
  if(!note) return;
  const d = await api('/api/cases/'+currentCaseId);
  const tl = (d.case.timeline||[]).concat([{date, note}]);
  await api('/api/cases/'+currentCaseId, {method:'PUT', body: JSON.stringify({timeline: tl})});
  document.getElementById('cd-tl-date').value=''; document.getElementById('cd-tl-note').value='';
  renderTimeline(tl);
}
async function updateCaseStatus(){
  const status = document.getElementById('cd-status-select').value;
  await api('/api/cases/'+currentCaseId, {method:'PUT', body: JSON.stringify({status})});
  await refreshCases(); renderStats(); renderActiveCases(cases); renderCaseList(cases);
  toast('Case status updated.');
}
async function deleteCase(){
  if(!confirm('Delete this case file? This cannot be undone.')) return;
  await api('/api/cases/'+currentCaseId, {method:'DELETE'});
  closeModal('modal-casedetail');
  await refreshCases(); renderStats(); renderActiveCases(cases); renderCaseList(cases);
  toast('Case file deleted.');
}

// ── Boot ───────────────────────────────────────────────────────────────
document.getElementById('in-email') && document.getElementById('in-email').addEventListener('keydown', e=>{ if(e.key==='Enter') doLogin(); });
tryResumeSession();
</script>
</body>
</html>
"""

HTML = HTML_HEAD + HTML_BODY + HTML_SCRIPT


# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES — auth
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/auth/dev', methods=['POST'])
@app.route('/api/auth/login', methods=['POST'])
def simple_login():
    """Simple name + email login — works in all environments."""
    data  = request.json or {}
    email = data.get('email', '').strip().lower()
    name  = data.get('name', '').strip() or email.split('@')[0]
    if not email or '@' not in email:
        return jsonify({'success': False, 'message': 'Valid email required'}), 400
    user_id = 'u_' + email.replace('@', '_').replace('.', '_')
    with get_db() as db:
        user = db.execute('SELECT * FROM users WHERE email=?', (email,)).fetchone()
        if user:
            db.execute("UPDATE users SET name=?,last_login=datetime('now') WHERE email=?", (name, email))
            user_id = user['id']
        else:
            db.execute("INSERT INTO users (id,email,name,last_login) VALUES (?,?,?,datetime('now'))",
                       (user_id, email, name))
    tok = secrets.token_urlsafe(32)
    session_set(tok, email)
    sessions[tok]['user_id'] = user_id
    sessions[tok]['name'] = name
    return jsonify({'success': True, 'token': tok, 'email': email, 'name': name})


@app.route('/api/auth/logout', methods=['POST'])
def logout():
    tok = request.headers.get('Authorization', '').replace('Bearer ', '')
    session_delete(tok)
    return jsonify({'success': True})


# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES — case file system
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/cases', methods=['GET', 'POST'])
def cases_collection():
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if request.method == 'POST':
        data = request.json or {}
        cid = uuid.uuid4().hex[:12]
        with get_db() as db:
            db.execute("""INSERT INTO case_files
                (id, owner_email, case_no, court, bench, parties, advocate, filing_date, doc_type, status, timeline)
                VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                (cid, sess['email'], data.get('case_no',''), data.get('court',''), data.get('bench',''),
                 data.get('parties',''), data.get('advocate',''), data.get('filing_date',''),
                 data.get('doc_type',''), data.get('status','Drafting'),
                 json.dumps(data.get('timeline', []))))
        return jsonify({'success': True, 'id': cid})

    q = (request.args.get('q') or '').strip().lower()
    with get_db() as db:
        rows = db.execute('SELECT * FROM case_files WHERE owner_email=? ORDER BY updated_at DESC',
                           (sess['email'],)).fetchall()
    cases = [row2dict(r) for r in rows]
    for c in cases:
        try: c['timeline'] = json.loads(c.get('timeline') or '[]')
        except Exception: c['timeline'] = []
    if q:
        def hit(c):
            hay = ' '.join(str(c.get(k,'')) for k in
                            ('case_no','court','bench','parties','advocate','doc_type')).lower()
            return q in hay
        cases = [c for c in cases if hit(c)]
    return jsonify({'success': True, 'cases': cases})


@app.route('/api/cases/<cid>', methods=['GET', 'PUT', 'DELETE'])
def case_item(cid):
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    with get_db() as db:
        row = db.execute('SELECT * FROM case_files WHERE id=?', (cid,)).fetchone()
        if not row or (row['owner_email'] != sess['email'] and not is_admin(sess)):
            return jsonify({'success': False, 'message': 'Not found'}), 404

        if request.method == 'DELETE':
            db.execute('DELETE FROM case_files WHERE id=?', (cid,))
            return jsonify({'success': True})

        if request.method == 'PUT':
            data = request.json or {}
            fields = ['case_no','court','bench','parties','advocate','filing_date','doc_type','status']
            sets = ', '.join(f'{f}=?' for f in fields)
            vals = [data.get(f, row[f]) for f in fields]
            timeline = data.get('timeline')
            if timeline is not None:
                sets += ', timeline=?'
                vals.append(json.dumps(timeline))
            sets += ", updated_at=datetime('now')"
            vals.append(cid)
            db.execute(f'UPDATE case_files SET {sets} WHERE id=?', vals)
            row = db.execute('SELECT * FROM case_files WHERE id=?', (cid,)).fetchone()

        drafts = db.execute('SELECT id,title,doc_type,status,updated_at FROM drafts WHERE case_id=? ORDER BY updated_at DESC',
                             (cid,)).fetchall()
    c = row2dict(row)
    try: c['timeline'] = json.loads(c.get('timeline') or '[]')
    except Exception: c['timeline'] = []
    c['drafts'] = [row2dict(d) for d in drafts]
    return jsonify({'success': True, 'case': c})


# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES — drafts, versions, comments
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/drafts', methods=['GET', 'POST'])
def drafts_collection():
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if request.method == 'POST':
        data = request.json or {}
        did = uuid.uuid4().hex[:12]
        with get_db() as db:
            db.execute("""INSERT INTO drafts (id, owner_email, case_id, title, doc_type, status, content)
                VALUES (?,?,?,?,?,?,?)""",
                (did, sess['email'], data.get('case_id'), data.get('title','Untitled Draft'),
                 data.get('doc_type',''), data.get('status','Drafting'), data.get('content','')))
        return jsonify({'success': True, 'id': did})

    q = (request.args.get('q') or '').strip().lower()
    with get_db() as db:
        rows = db.execute('SELECT * FROM drafts WHERE owner_email=? ORDER BY updated_at DESC',
                           (sess['email'],)).fetchall()
    drafts = [row2dict(r) for r in rows]
    if q:
        def hit(d):
            hay = ' '.join(str(d.get(k,'')) for k in ('title','doc_type')).lower()
            return q in hay
        drafts = [d for d in drafts if hit(d)]
    return jsonify({'success': True, 'drafts': drafts})


@app.route('/api/drafts/<did>', methods=['GET', 'PUT', 'DELETE'])
def draft_item(did):
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    with get_db() as db:
        row = db.execute('SELECT * FROM drafts WHERE id=?', (did,)).fetchone()
        if not row or (row['owner_email'] != sess['email'] and not is_admin(sess)):
            return jsonify({'success': False, 'message': 'Not found'}), 404

        if request.method == 'DELETE':
            db.execute('DELETE FROM drafts WHERE id=?', (did,))
            db.execute('DELETE FROM draft_versions WHERE draft_id=?', (did,))
            db.execute('DELETE FROM draft_comments WHERE draft_id=?', (did,))
            return jsonify({'success': True})

        if request.method == 'PUT':
            data = request.json or {}
            fields = ['title','doc_type','status','content','case_id']
            sets = ', '.join(f'{f}=?' for f in fields) + ", updated_at=datetime('now')"
            vals = [data.get(f, row[f]) for f in fields] + [did]
            db.execute(f'UPDATE drafts SET {sets} WHERE id=?', vals)
            if data.get('save_version'):
                vid = uuid.uuid4().hex[:12]
                db.execute('INSERT INTO draft_versions (id, draft_id, content, note) VALUES (?,?,?,?)',
                           (vid, did, data.get('content',''), data.get('version_note','')))
            row = db.execute('SELECT * FROM drafts WHERE id=?', (did,)).fetchone()

        versions = db.execute('SELECT id,note,created_at FROM draft_versions WHERE draft_id=? ORDER BY created_at DESC',
                               (did,)).fetchall()
        comments = db.execute('SELECT id,author,text,created_at FROM draft_comments WHERE draft_id=? ORDER BY created_at ASC',
                               (did,)).fetchall()
    d = row2dict(row)
    d['versions'] = [row2dict(v) for v in versions]
    d['comments'] = [row2dict(c) for c in comments]
    return jsonify({'success': True, 'draft': d})


@app.route('/api/drafts/<did>/versions/<vid>')
def draft_version(did, vid):
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    with get_db() as db:
        d = db.execute('SELECT owner_email FROM drafts WHERE id=?', (did,)).fetchone()
        if not d or (d['owner_email'] != sess['email'] and not is_admin(sess)):
            return jsonify({'success': False, 'message': 'Not found'}), 404
        v = db.execute('SELECT * FROM draft_versions WHERE id=? AND draft_id=?', (vid, did)).fetchone()
        if not v:
            return jsonify({'success': False, 'message': 'Version not found'}), 404
    return jsonify({'success': True, 'version': row2dict(v)})


@app.route('/api/drafts/<did>/comments', methods=['POST'])
def add_comment(did):
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    data = request.json or {}
    text = (data.get('text') or '').strip()
    if not text:
        return jsonify({'success': False, 'message': 'Comment text required'}), 400
    with get_db() as db:
        d = db.execute('SELECT owner_email FROM drafts WHERE id=?', (did,)).fetchone()
        if not d or (d['owner_email'] != sess['email'] and not is_admin(sess)):
            return jsonify({'success': False, 'message': 'Not found'}), 404
        cid = uuid.uuid4().hex[:12]
        author = sess.get('name') or sess.get('email')
        db.execute('INSERT INTO draft_comments (id, draft_id, author, text) VALUES (?,?,?,?)',
                   (cid, did, author, text))
    return jsonify({'success': True, 'id': cid})


# ═══════════════════════════════════════════════════════════════════════════════
#  ROUTES — AI drafting + intelligence
# ═══════════════════════════════════════════════════════════════════════════════

@app.route('/api/legal/generate', methods=['POST'])
def gen_ai_legal_draft():
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if not os.environ.get('GROQ_API_KEY', '').strip():
        return jsonify({'success': False,
                        'message': 'GROQ_API_KEY not set. Get a free key at https://console.groq.com'}), 400

    is_multipart = request.content_type and 'multipart/form-data' in request.content_type
    mode = (request.form.get('mode') if is_multipart else (request.json or {}).get('mode')) or 'custom'

    try:
        if mode == 'format':
            details = (request.form.get('details') or '').strip()
            if not details:
                return jsonify({'success': False, 'message': 'Please provide the data to fill into the format.'}), 400
            file_storage = request.files.get('format_file')
            if not file_storage or not file_storage.filename:
                return jsonify({'success': False, 'message': 'Please upload a format/sample document.'}), 400
            reference_text = extract_text_from_upload(file_storage)
            doc_type = 'Legal Draft'
            ai_text = ai_draft_legal_document(doc_type, details, reference_text=reference_text)
            draft_id = request.form.get('draft_id')
        else:
            data = request.json or {}
            doc_type = (data.get('doc_type') or '').strip()
            details  = (data.get('details') or '').strip()
            if not doc_type or not details:
                return jsonify({'success': False, 'message': 'Please provide the document type and details.'}), 400
            ai_text = ai_draft_legal_document(doc_type, details)
            draft_id = data.get('draft_id')

        path = build_ai_legal_docx(doc_type, ai_text)
        jid  = uuid.uuid4().hex
        jobs[jid] = {'file_path': path, 'topic': doc_type, 'owner_email': sess['email'], 'preview': ai_text}

        # If tied to a draft, persist the generated text as the draft content + a version snapshot
        if draft_id:
            with get_db() as db:
                row = db.execute('SELECT * FROM drafts WHERE id=?', (draft_id,)).fetchone()
                if row and row['owner_email'] == sess['email']:
                    db.execute("UPDATE drafts SET content=?, doc_type=?, updated_at=datetime('now') WHERE id=?",
                               (ai_text, doc_type, draft_id))
                    vid = uuid.uuid4().hex[:12]
                    db.execute('INSERT INTO draft_versions (id, draft_id, content, note) VALUES (?,?,?,?)',
                               (vid, draft_id, ai_text, 'AI draft generated'))

        return jsonify({'success': True, 'job_id': jid, 'preview': ai_text, 'doc_type': doc_type})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/legal/assist', methods=['POST'])
def legal_assist():
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    if not os.environ.get('GROQ_API_KEY', '').strip():
        return jsonify({'success': False,
                        'message': 'GROQ_API_KEY not set. Get a free key at https://console.groq.com'}), 400
    data = request.json or {}
    action = (data.get('action') or '').strip()
    text   = data.get('text') or ''
    extra  = data.get('extra') or ''
    try:
        result = ai_assist(action, text, extra)
        return jsonify({'success': True, 'result': result, 'action': action})
    except ValueError as e:
        return jsonify({'success': False, 'message': str(e)}), 400
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/legal/actions')
def legal_actions():
    return jsonify({'success': True, 'actions': [
        {'id': k, 'label': v['label']} for k, v in ASSIST_ACTIONS.items()
    ]})


@app.route('/api/download/<jid>')
def download_draft(jid):
    sess = require_session()
    if not sess:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    job = jobs.get(jid)
    if not job:
        return jsonify({'success': False, 'message': 'Job not found'}), 404
    if job['owner_email'] != sess['email'] and not is_admin(sess):
        return jsonify({'success': False, 'message': 'Forbidden'}), 403

    fp = job.get('file_path')
    if not fp or not os.path.exists(fp):
        return jsonify({'success': False, 'message': 'File not found on server'}), 404

    slug = re.sub(r'[^\w\-]', '_', (job.get('topic') or jid)[:40])
    return send_file(fp, as_attachment=True,
                     download_name=f'dratido_{slug}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/')
def index():
    return Response(HTML, mimetype='text/html')


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    os.makedirs('generated', exist_ok=True)

    groq_key = os.environ.get('GROQ_API_KEY', '').strip()
    key_str = '\u2713 Groq \u2014 ready!' if groq_key else '\u2717 NOT SET \u2014 see below'
    print('\n' + '=' * 60)
    print(f'  {APP_NAME} \u2014 {APP_TAGLINE}')
    print('  AI Legal Drafting for the Indian court office')
    print('  Powered by Groq (free tier)')
    print('  Open browser:  http://127.0.0.1:8081')
    print(f'  GROQ_API_KEY: {key_str}')
    print('=' * 60 + '\n')
    if not groq_key:
        print('  Get your free Groq API key at https://console.groq.com')
        print('  then:  export GROQ_API_KEY=your_key_here   (Mac/Linux)')
        print('         set GROQ_API_KEY=your_key_here       (Windows)\n')

    port = int(os.environ.get('PORT', 8081))
    app.run(host='0.0.0.0', port=port, debug=False, threaded=True)

